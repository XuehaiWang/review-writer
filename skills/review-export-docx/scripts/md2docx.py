#!/usr/bin/env python3
"""
md2docx.py -- Convert a Markdown file to DOCX using review_template.docx styles.

Inline support:
  **bold**  *italic*  ***bold-italic***  `code`
  ^superscript^       alnum_subscript_   $math$  $$display math$$
  [@citation]  ->  [citation key]

Section-aware styling (auto-detected from headings OR bold-only paragraphs):
  Abstract / Keywords / Acknowledgments / References / Supporting Information

Font specification (explicitly applied to every run):
  H1 title      : Times New Roman 18 pt  (xiao-er)
  Author line   : Times New Roman 12 pt  (xiao-si)
  Affiliation   : Times New Roman 10.5 pt (wu-hao)
  H2 heading    : Times New Roman 14 pt  (si-hao)  bold
  H3 heading    : Times New Roman 12 pt  bold italic
  Body / Abstract / Keywords : Times New Roman 12 pt
  Captions / References      : Times New Roman 10.5 pt

Usage:
    python3 scripts/md2docx.py --input review.md --output review.docx
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
from copy import deepcopy  # noqa: F401
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from PIL import Image as PILImage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml  # noqa: F401
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    from latex2word import LatexToWordElement
    _LATEX_OK = True
except ImportError:
    _LATEX_OK = False

# ---------------------------------------------------------------------------
# Template style name map
# ---------------------------------------------------------------------------

_S: Dict[str, str] = {
    "title":        "BA_Title",
    "author":       "BB_Author_Name",
    "address":      "BC_Author_Address",
    "email":        "BI_Email_Address",
    "abstract":     "BD_Abstract",
    "keywords":     "BG_Keywords",
    "body":         "TA_Main_Text",
    "figure":       "VA_Figure_Caption",
    "table_title":  "VD_Table_Title",
    "table_body":   "TC_Table_Body",
    "chart":        "VB_Chart_Title",
    "scheme":       "VC_Scheme_Title",
    "references":   "TF_References_Section",
    "acks":         "TD_Acknowledgments",
    "supporting":   "TE_Supporting_Information",
    "footnote":     "FA_Corresponding_Author_Footnote",
}

# ---------------------------------------------------------------------------
# Font spec -- every run gets an explicit font name + size
# ---------------------------------------------------------------------------

_FONT_SPEC: Dict[str, Dict] = {
    "title":        {"font": "Times New Roman", "size": 18},
    "author":       {"font": "Times New Roman", "size": 12},
    "address":      {"font": "Times New Roman", "size": 10.5},
    "email":        {"font": "Times New Roman", "size": 10.5},
    "abstract":     {"font": "Times New Roman", "size": 12},
    "keywords":     {"font": "Times New Roman", "size": 12},
    "body":         {"font": "Times New Roman", "size": 12},
    "h2":           {"font": "Times New Roman", "size": 14,  "bold": True},
    "h3":           {"font": "Times New Roman", "size": 12,  "bold": True, "italic": True},
    "h4":           {"font": "Times New Roman", "size": 12,  "italic": True},
    "figure":       {"font": "Times New Roman", "size": 10.5},
    "table_title":  {"font": "Times New Roman", "size": 12},
    "table_body":   {"font": "Times New Roman", "size": 10.5},
    "scheme":       {"font": "Times New Roman", "size": 10.5},
    "chart":        {"font": "Times New Roman", "size": 10.5},
    "references":   {"font": "Times New Roman", "size": 10.5},
    "acks":         {"font": "Times New Roman", "size": 12},
    "supporting":   {"font": "Times New Roman", "size": 12},
    "footnote":     {"font": "Times New Roman", "size": 10.5},
}

# Heading level -> (para_style_key, font_spec_key)
_HEADING_FORMAT: Dict[int, Tuple[str, str]] = {
    1: ("title", "title"),
    2: ("body",  "h2"),
    3: ("body",  "h3"),
    4: ("body",  "h4"),
    5: ("body",  "body"),
    6: ("body",  "body"),
}

_SECTION_CONTEXT: Dict[str, str] = {
    "abstract":               "abstract",
    "keywords":               "keywords",
    "key words":              "keywords",
    "acknowledgments":        "acks",
    "acknowledgements":       "acks",
    "supporting information": "supporting",
    "references":             "references",
    "reference":              "references",
}

_INTRODUCTION_HEADINGS = {"introduction", "background"}

_CAPTION_PATTERNS: List[Tuple[re.Pattern, str]] = [
    (re.compile(r"^(figure|fig\.)\s*\d+", re.I), "figure"),
    (re.compile(r"^table\s*\d+",            re.I), "table_title"),
    (re.compile(r"^scheme\s*\d+",           re.I), "scheme"),
    (re.compile(r"^chart\s*\d+",            re.I), "chart"),
]

# MinerU sometimes emits LaTex without a $...$ delimiter and with spaces
# between commands, braces, and chemical symbols. These patterns cover the
# chemistry notation found in the review corpus while deliberately leaving
# Markdown image paths (for example figures\\figure_01.png) untouched.
_RAW_LATEX_REWRITES: List[Tuple[re.Pattern, Any]] = [
    (
        re.compile(r"\{\s*\\mathrm\s*\{\s*([^{}]+?)\s*\}\s*\}\s*_\s*\{\s*([^{}]+?)\s*\}"),
        lambda m: "$\\mathrm{" + re.sub(r"\s+", "", m.group(1)) + "}_{" + re.sub(r"\s+", "", m.group(2)) + "}$",
    ),
    (
        re.compile(r"\{\s*\\pmb\s*\{\s*\\(alpha|beta|gamma|delta|epsilon|theta|lambda|mu|nu|pi|rho|sigma|tau|phi|chi|psi|omega)\s*\}\s*\}\s*\^\s*\{\s*\\prime\s*\}"),
        lambda m: "$\\" + m.group(1) + "^\\prime$",
    ),
    (
        re.compile(r"\{\s*\\pmb\s*\{\s*\\(alpha|beta|gamma|delta|epsilon|theta|lambda|mu|nu|pi|rho|sigma|tau|phi|chi|psi|omega)\s*\}\s*\}"),
        lambda m: "$\\" + m.group(1) + "$",
    ),
    (
        re.compile(r"\b([A-Za-z])\s*_\s*\{\s*([A-Za-z])\s*\^\s*\{\s*(\d+)\s*\}\s*\}\s*\\prime"),
        lambda m: f"${m.group(1)}_{{{m.group(2)}^{m.group(3)}}}\\prime$",
    ),
]
_RAW_GREEK_COMMAND_RE = re.compile(
    r"\\(alpha|beta|gamma|delta|epsilon|theta|lambda|mu|nu|pi|rho|sigma|tau|phi|chi|psi|omega)\b"
)
_RAW_LATEX_COMMAND_RE = re.compile(r"\\(?!figure_[A-Za-z0-9_-]+\.[A-Za-z0-9]+)[A-Za-z]+")

_TRUNCATED_GREEK_ESCAPE_RE = re.compile(r"\x03([0-9a-fA-F]{2})")
_KNOWN_TRUNCATED_UNICODE = {
    "\x02": "\u2032",  # U+2032 PRIME
    "\x13": "\u2013",  # U+2013 EN DASH
    "\x14": "\u2014",  # U+2014 EM DASH
}


def make_xml_compatible(text: str) -> Tuple[str, int]:
    """Recover known relay damage and replace remaining XML-invalid chars."""
    repaired = re.sub(r"(?<=C)\x03(?=C)", "\u2013", str(text or ""))
    repaired = _TRUNCATED_GREEK_ESCAPE_RE.sub(
        lambda match: chr(int("03" + match.group(1), 16)),
        repaired,
    )
    repaired = "".join(_KNOWN_TRUNCATED_UNICODE.get(char, char) for char in repaired)
    output: List[str] = []
    replaced = 0
    for char in repaired:
        number = ord(char)
        if (
            number in {0x09, 0x0A, 0x0D}
            or 0x20 <= number <= 0xD7FF
            or 0xE000 <= number <= 0xFFFD
            or 0x10000 <= number <= 0x10FFFF
        ):
            output.append(char)
        else:
            output.append("\uFFFD")
            replaced += 1
    return "".join(output), replaced


def normalize_mineru_latex(md_text: str) -> str:
    """Turn known raw MinerU LaTex into Math-delimited expressions for DOCX."""
    for pattern, replacement in _RAW_LATEX_REWRITES:
        md_text = pattern.sub(replacement, md_text)

    # Do not use a simple lookbehind for Greek commands: a formula such as
    # ``$\\alpha, \\alpha^\\prime$`` contains a later command that is not
    # immediately adjacent to its opening ``$``.  Re-wrapping that command
    # would split one valid formula into mismatched math delimiters.
    source = md_text
    md_text = _RAW_GREEK_COMMAND_RE.sub(
        lambda match: (
            match.group(0)
            if source[:match.start()].count("$") % 2
            else "$" + match.group(0) + "$"
        ),
        source,
    )
    leftovers = [
        match.group(0)
        for match in _RAW_LATEX_COMMAND_RE.finditer(md_text)
        if md_text[:match.start()].count("$") % 2 == 0
    ]
    if leftovers:
        examples = ", ".join(dict.fromkeys(leftovers[:5]))
        raise ValueError(f"Unconverted MinerU LaTex remains in Markdown: {examples}")
    return md_text


def _usable_page_width_inches(doc: Document) -> float:
    section = doc.sections[0]
    width_emu = section.page_width - section.left_margin - section.right_margin
    # 914400 EMUs per inch. Keep a conservative upper bound for journal templates.
    return max(1.0, min(6.2, width_emu / 914400))

_UNICODE_SUPERSCRIPT_MAP: Dict[str, str] = {
    "⁰": "0",
    "¹": "1",
    "²": "2",
    "³": "3",
    "⁴": "4",
    "⁵": "5",
    "⁶": "6",
    "⁷": "7",
    "⁸": "8",
    "⁹": "9",
    "⁺": "+",
    "⁻": "-",
    "⁼": "=",
    "⁽": "(",
    "⁾": ")",
}

_UNICODE_SUBSCRIPT_MAP: Dict[str, str] = {
    "₀": "0",
    "₁": "1",
    "₂": "2",
    "₃": "3",
    "₄": "4",
    "₅": "5",
    "₆": "6",
    "₇": "7",
    "₈": "8",
    "₉": "9",
    "₊": "+",
    "₋": "-",
    "₌": "=",
    "₍": "(",
    "₎": ")",
}

# ---------------------------------------------------------------------------
# Run dataclass
# ---------------------------------------------------------------------------

@dataclass
class Run:
    text:        str  = ""
    bold:        bool = False
    italic:      bool = False
    code:        bool = False
    superscript: bool = False
    subscript:   bool = False
    math:        str  = ""

# ---------------------------------------------------------------------------
# Inline parser
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r"(\$\$[\s\S]*?\$\$"
    r"|\$[^$\n]+?\$"
    r"|\*\*\*(?:\S[^*\n]*?\S|\S)\*\*\*"
    r"|\*\*(?:\S[^*\n]*?\S|\S)\*\*"
    r"|\^[^\^\s\n]+?\^"
    r"|_[^_\s\n]+_"
    r"|\*(?:\S[^*\n]*?\S|\S)\*"
    r"|__(?:\S[^_\n]*?\S|\S)__"
    r"|_(?:\S[^_\n]*?\S|\S)_"
    r"|`[^`\n]+`"
    r"|\[@([^\]]+)\]"
    r"|\[([^\]]*)\]\([^\)]*\)"
    r")"
)


def _parse_nested(inner: str, bold: bool = False, italic: bool = False) -> List[Run]:
    runs = parse_inline(inner)
    for r in runs:
        if bold:
            r.bold = True
        if italic:
            r.italic = True
    return runs


def parse_inline(raw: str) -> List[Run]:
    runs: List[Run] = []
    pos = 0
    for m in _INLINE_RE.finditer(raw):
        if m.start() > pos:
            runs.append(Run(text=raw[pos:m.start()]))
        token = m.group(0)
        char_before = raw[m.start() - 1] if m.start() > 0 else ""

        if token.startswith("$$"):
            runs.append(Run(math=token[2:-2].strip()))
        elif token.startswith("$"):
            runs.append(Run(math=token[1:-1].strip()))
        elif token.startswith("***"):
            runs.extend(_parse_nested(token[3:-3], bold=True, italic=True))
        elif token.startswith("**"):
            runs.extend(_parse_nested(token[2:-2], bold=True))
        elif token.startswith("^") and token.endswith("^"):
            runs.append(Run(text=token[1:-1], superscript=True))
        elif token.startswith("_") and token.endswith("_") and " " not in token[1:-1]:
            if char_before.isalnum():
                runs.append(Run(text=token[1:-1], subscript=True))
            else:
                runs.append(Run(text=token[1:-1], italic=True))
        elif token.startswith("*"):
            runs.extend(_parse_nested(token[1:-1], italic=True))
        elif token.startswith("__"):
            runs.extend(_parse_nested(token[2:-2], bold=True))
        elif token.startswith("_"):
            runs.extend(_parse_nested(token[1:-1], italic=True))
        elif token.startswith("`"):
            runs.append(Run(text=token[1:-1], code=True))
        elif token.startswith("[@"):
            cite_key = m.group(1) or token[2:-1]
            runs.append(Run(text=f"[{cite_key}]"))
        elif token.startswith("["):
            display = m.group(2)
            runs.append(Run(text=display if display is not None else token))
        else:
            runs.append(Run(text=token))
        pos = m.end()

    if pos < len(raw):
        runs.append(Run(text=raw[pos:]))
    return runs or [Run(text=raw)]


# ---------------------------------------------------------------------------
# Font + run application
# ---------------------------------------------------------------------------

_LATEX_SYMBOLS = {
    "alpha": "\u03b1", "beta": "\u03b2", "gamma": "\u03b3", "delta": "\u03b4",
    "epsilon": "\u03b5", "theta": "\u03b8", "lambda": "\u03bb", "mu": "\u03bc",
    "nu": "\u03bd", "pi": "\u03c0", "rho": "\u03c1", "sigma": "\u03c3",
    "tau": "\u03c4", "phi": "\u03c6", "chi": "\u03c7", "psi": "\u03c8",
    "omega": "\u03c9", "Gamma": "\u0393", "Delta": "\u0394", "Theta": "\u0398",
    "Lambda": "\u039b", "Xi": "\u039e", "Pi": "\u03a0", "Sigma": "\u03a3",
    "Phi": "\u03a6", "Psi": "\u03a8", "Omega": "\u03a9", "times": "\u00d7",
    "cdot": "\u00b7", "pm": "\u00b1", "leq": "\u2264", "geq": "\u2265",
    "neq": "\u2260", "approx": "\u2248", "rightarrow": "\u2192", "leftarrow": "\u2190",
    "leftrightarrow": "\u2194", "sum": "\u03a3", "prod": "\u03a0", "infty": "\u221e",
}
_LATEX_TEXT_WRAPPERS = "mathrm|mathbf|mathit|text|operatorname|pmb|boldsymbol|rm|bf|it"


def _latex_to_readable_text(latex: str) -> str:
    """Remove unsupported LaTex syntax while retaining the displayed meaning."""
    text = latex.strip()
    # Repeat so adjacent, non-nested formatting wrappers are all unwrapped.
    wrapper_re = re.compile(r"\\(?:" + _LATEX_TEXT_WRAPPERS + r")\s*\{([^{}]*)\}")
    previous = None
    while previous != text:
        previous = text
        text = wrapper_re.sub(lambda m: re.sub(r"\s+", "", m.group(1)) if m.group(0).startswith("\\mathrm") else m.group(1), text)
        text = re.sub(r"\\(?:d?frac)\s*\{([^{}]*)\}\s*\{([^{}]*)\}", r"(\1)/(\2)", text)

    text = text.replace("\\prime", "\u2032")
    text = re.sub(r"\\([A-Za-z]+)", lambda m: _LATEX_SYMBOLS.get(m.group(1), m.group(1)), text)
    # TeX spacing commands are no longer meaningful once rendered as Word text.
    text = text.replace("\\,", "").replace("\\;", " ").replace("\\!", "")
    return text


def _extract_script_content(text: str, start: int) -> Tuple[str, int]:
    """Return a LaTex script token and the index immediately after it."""
    if start >= len(text):
        return "", start
    if text[start] != "{":
        return text[start], start + 1
    depth = 1
    index = start + 1
    while index < len(text) and depth:
        if text[index] == "{":
            depth += 1
        elif text[index] == "}":
            depth -= 1
        index += 1
    return text[start + 1:index - 1] if depth == 0 else text[start + 1:], index


def _split_latex_script_segments(text: str, default_mode: str = "normal") -> List[Tuple[str, str]]:
    """Preserve LaTex _{...} and ^{...} using Word-native run formatting."""
    segments: List[Tuple[str, str]] = []
    buffer = ""

    def append(mode: str, value: str) -> None:
        if not value:
            return
        if segments and segments[-1][0] == mode:
            segments[-1] = (mode, segments[-1][1] + value)
        else:
            segments.append((mode, value))

    index = 0
    while index < len(text):
        char = text[index]
        if char in "_^":
            append(default_mode, buffer)
            buffer = ""
            script, index = _extract_script_content(text, index + 1)
            script_mode = "subscript" if char == "_" else "superscript"
            for mode, value in _split_latex_script_segments(script, script_mode):
                append(mode, value)
        elif char in "{}":
            index += 1
        else:
            buffer += char
            index += 1
    append(default_mode, buffer)
    return segments


def _apply_math(para, latex: str) -> None:
    if _LATEX_OK:
        try:
            LatexToWordElement(latex).add_latex_to_paragraph(para)
            return
        except Exception:
            pass
    # Do not leak raw LaTex when the optional OMML converter is unavailable.
    # The fallback keeps formulas readable and carries their scripts as native
    # Word run formatting.
    readable = _latex_to_readable_text(latex)
    for segment_mode, segment_text in _split_latex_script_segments(readable):
        run = para.add_run(segment_text)
        run.italic = True
        if segment_mode == "superscript":
            run.font.superscript = True
        elif segment_mode == "subscript":
            run.font.subscript = True


def _split_script_segments(text: str) -> List[Tuple[str, str]]:
    segments: List[Tuple[str, str]] = []
    mode = "normal"
    current = ""

    def looks_like_ascii_subscript(index: int) -> bool:
        if index < 0 or index >= len(text):
            return False
        char = text[index]
        if not char.isdigit() or index == 0:
            return False
        prev = text[index - 1]
        if prev in "-–—/[ ":
            return False
        if prev.isalpha():
            return True
        if prev in ")]}" and index >= 2 and text[index - 2].isalpha():
            return True
        return False

    def flush() -> None:
        nonlocal current
        if current:
            segments.append((mode, current))
            current = ""

    for idx, char in enumerate(text):
        if char in _UNICODE_SUPERSCRIPT_MAP:
            char_mode = "superscript"
            rendered = _UNICODE_SUPERSCRIPT_MAP[char]
        elif char in _UNICODE_SUBSCRIPT_MAP:
            char_mode = "subscript"
            rendered = _UNICODE_SUBSCRIPT_MAP[char]
        elif looks_like_ascii_subscript(idx):
            char_mode = "subscript"
            rendered = char
        else:
            char_mode = "normal"
            rendered = char
        if char_mode != mode:
            flush()
            mode = char_mode
        current += rendered
    flush()
    return segments or [("normal", text)]


def apply_runs(
    para,
    runs: List[Run],
    spec_key: str = "body",
    force_bold: bool = False,
    force_italic: bool = False,
) -> None:
    spec        = _FONT_SPEC.get(spec_key, _FONT_SPEC["body"])
    font_name   = spec["font"]
    size_pt     = spec["size"]
    spec_bold   = spec.get("bold", False)
    spec_italic = spec.get("italic", False)

    for r in runs:
        if r.math:
            _apply_math(para, r.math)
            continue
        segments = [("normal", r.text)] if r.code else _split_script_segments(r.text)
        for segment_mode, segment_text in segments:
            if not segment_text:
                continue
            wr = para.add_run(segment_text)
            if r.code:
                wr.font.name = "Courier New"
                wr.font.size = Pt(9)
            else:
                wr.font.name = font_name
                wr.font.size = Pt(size_pt)
            wr.bold = (spec_bold or force_bold or r.bold) or None
            wr.italic = (spec_italic or force_italic or r.italic) or None
            if r.superscript or segment_mode == "superscript":
                wr.font.superscript = True
            if r.subscript or segment_mode == "subscript":
                wr.font.subscript = True


# ---------------------------------------------------------------------------
# Paragraph factory
# ---------------------------------------------------------------------------

def _para(
    doc: Document,
    style_key: str,
    spec_key: str,
    inline_text: str = "",
    force_bold: bool = False,
    force_italic: bool = False,
):
    p = doc.add_paragraph(style=_S.get(style_key, _S["body"]))
    if inline_text:
        apply_runs(p, parse_inline(inline_text),
                   spec_key=spec_key,
                   force_bold=force_bold,
                   force_italic=force_italic)
    return p


# ---------------------------------------------------------------------------
# Table builder
# ---------------------------------------------------------------------------

def _set_cell_borders(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        tcPr.append(elem)


def _add_table(doc: Document, header: List[str], rows: List[List[str]]) -> None:
    ncols = max(len(header), max((len(r) for r in rows), default=1))
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        cell.text = ""
        cell.paragraphs[0].style = doc.styles[_S["table_body"]]
        apply_runs(cell.paragraphs[0], parse_inline(h),
                   spec_key="table_body", force_bold=True)
        _set_cell_borders(cell)
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i + 1, j)
            cell.text = ""
            cell.paragraphs[0].style = doc.styles[_S["table_body"]]
            apply_runs(cell.paragraphs[0],
                       parse_inline(row[j] if j < len(row) else ""),
                       spec_key="table_body")
            _set_cell_borders(cell)


# ---------------------------------------------------------------------------
# Block tokenizer
# ---------------------------------------------------------------------------

@dataclass
class Block:
    kind:     str
    level:    int             = 0
    text:     str             = ""
    ordered:  bool            = False
    depth:    int             = 0
    code:     str             = ""
    language: str             = ""
    header:   List[str]       = field(default_factory=list)
    rows:     List[List[str]] = field(default_factory=list)
    alt:      str             = ""
    path:     str             = ""
    latex:    str             = ""
    lines:    List[str]       = field(default_factory=list)


@dataclass
class SummaryChartBundle:
    full: Path
    sections: Dict[str, Tuple[str, Path]]


_HEADING_RE    = re.compile(r"^ {0,3}(#{1,6})\s+(.*)")
_EMBEDDED_HEADING_PREFIX_RE = re.compile(r"^#{1,6}\s+")
_NUMBERED_SECTION_HEADING_RE = re.compile(r"^\s*\d+(?:\.\d+)*[.)]?\s+\S")
_HTML_ANCHOR_RE = re.compile(r"^<a\s+id=[\"']ref-\d+[\"']\s*>\s*</a>\s*$", re.I)
_PARAGRAPH_ID_MARKER_RE = re.compile(r"^\s*<!--\s*paragraph_id:\s*[A-Za-z0-9_.:-]+\s*-->\s*$")
_UL_RE         = re.compile(r"^(\s*)[-*+]\s+(.*)")
_OL_RE         = re.compile(r"^(\s*)\d+[.)]\s+(.*)")
_FENCE_RE      = re.compile(r"^\s*(`{3,}|~{3,})(.*)$")
_MATH_FENCE_RE = re.compile(r"^\$\$\s*$")
_IMG_RE        = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
_TABLE_ROW_RE  = re.compile(r"^\|.+")
_HR_RE         = re.compile(r"^(?:-{3,}|_{3,}|\*{3,})\s*$")
_REF_ENTRY_RE  = re.compile(r"^\[?\d+\]?[.)\s]|\[@[^\]]+\]:")
_AFFIL_START   = re.compile(r"^\^[0-9,]+\^")
_INDENTED_RE   = re.compile(r"^(?: {4,}|\t+)(.*)$")


def _is_continuation(line: str) -> bool:
    if not line.strip():
        return False
    if _REF_ENTRY_RE.match(line):
        return False
    if _AFFIL_START.match(line):
        return False
    for pat in (_HEADING_RE, _FENCE_RE, _TABLE_ROW_RE,
                _UL_RE, _OL_RE, _HR_RE, _IMG_RE):
        if pat.match(line):
            return False
    return True


def tokenize(md_text: str) -> List[Block]:
    lines = md_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: List[Block] = []
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]

        # YAML front matter
        if i == 0 and line.strip() == "---":
            i += 1
            while i < n and lines[i].strip() != "---":
                i += 1
            i += 1
            continue

        # Fenced code block
        m = _FENCE_RE.match(line)
        if m:
            marker = m.group(1)
            language_info = m.group(2).strip()
            lang = language_info.split(maxsplit=1)[0] if language_info else ""
            closing_fence = re.compile(
                rf"^\s*{re.escape(marker[0])}{{{len(marker)},}}\s*$"
            )
            code_lines: List[str] = []
            i += 1
            while i < n and not closing_fence.match(lines[i]):
                code_lines.append(lines[i])
                i += 1
            if i < n:
                i += 1
            blocks.append(Block(kind="code_block", language=lang,
                                code="\n".join(code_lines)))
            continue

        # Display math fence
        if _MATH_FENCE_RE.match(line):
            math_lines: List[str] = []
            i += 1
            while i < n and not _MATH_FENCE_RE.match(lines[i]):
                math_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append(Block(kind="math_block", latex="\n".join(math_lines)))
            continue

        if _HTML_ANCHOR_RE.match(line.strip()):
            i += 1
            continue
        if _PARAGRAPH_ID_MARKER_RE.match(line):
            i += 1
            continue

        # ATX heading
        m = _HEADING_RE.match(line)
        if m:
            heading_text = m.group(2).strip()
            while _EMBEDDED_HEADING_PREFIX_RE.match(heading_text):
                heading_text = _EMBEDDED_HEADING_PREFIX_RE.sub("", heading_text, count=1).strip()
            blocks.append(Block(kind="heading", level=len(m.group(1)),
                                text=heading_text))
            i += 1
            continue

        # Horizontal rule
        if _HR_RE.match(line):
            blocks.append(Block(kind="hr"))
            i += 1
            continue

        # Standalone image
        m = _IMG_RE.match(line)
        if m:
            blocks.append(Block(kind="image", alt=m.group(1), path=m.group(2)))
            i += 1
            continue

        # Indented text block: preserve one source line -> one logical block line.
        m = _INDENTED_RE.match(line)
        if m and not _TABLE_ROW_RE.match(line):
            block_lines: List[str] = [m.group(1).rstrip()]
            i += 1
            while i < n:
                next_match = _INDENTED_RE.match(lines[i])
                if not next_match or not next_match.group(1).strip():
                    break
                block_lines.append(next_match.group(1).rstrip())
                i += 1
            blocks.append(Block(kind="indented_block", lines=block_lines))
            continue

        # Reference definition  [@key]: text...
        ref_m = re.match(r"^\[@([^\]]+)\]:\s*(.+)$", line)
        if ref_m:
            blocks.append(Block(kind="ref_def", text=ref_m.group(2).strip()))
            i += 1
            continue

        # Pipe table
        if _TABLE_ROW_RE.match(line):
            raw_rows: List[List[str]] = []
            while i < n and _TABLE_ROW_RE.match(lines[i]):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                raw_rows.append(cells)
                i += 1
            data = [r for r in raw_rows
                    if not all(re.match(r"^[-: ]+$", c) for c in r)]
            if data:
                blocks.append(Block(kind="table", header=data[0], rows=data[1:]))
            continue

        # Ordered list item
        m = _OL_RE.match(line)
        if m:
            blocks.append(Block(kind="list_item", ordered=True,
                                depth=len(m.group(1)) // 2, text=m.group(2)))
            i += 1
            continue

        # Unordered list item
        m = _UL_RE.match(line)
        if m:
            blocks.append(Block(kind="list_item", ordered=False,
                                depth=len(m.group(1)) // 2, text=m.group(2)))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Paragraph
        para_lines = [line]
        i += 1
        while i < n and _is_continuation(lines[i]):
            para_lines.append(lines[i].rstrip())
            i += 1
        blocks.append(Block(kind="paragraph",
                            text=" ".join(l.rstrip() for l in para_lines)))

    return blocks


# ---------------------------------------------------------------------------
# Context helpers
# ---------------------------------------------------------------------------

def _plain_text(raw: str) -> str:
    t = re.sub(r"\*+|__?", "", raw)
    t = re.sub(r"\[@[^\]]+\]", "", t)
    t = re.sub(r"\^[^\^]+\^", "", t)
    t = re.sub(r"`[^`]+`", "", t)
    t = re.sub(r"\[[^\]]*\]\([^\)]*\)", "", t)
    return t.strip()


def _section_ctx(text: str) -> Optional[str]:
    return _SECTION_CONTEXT.get(text.strip().lower())


def is_introduction_heading(text: str) -> bool:
    """Recognize the manuscript introduction that anchors the full-review chart."""
    title = re.sub(r"^\s*\d+(?:\.\d+)*[.)]?\s*", "", text).strip()
    return title.casefold() in _INTRODUCTION_HEADINGS


def should_insert_full_chart_before_heading(
    introduction_level: int | None, next_heading_level: int
) -> bool:
    """Place the chart after Introduction and before the next peer-level section."""
    return introduction_level is not None and next_heading_level <= introduction_level


def _caption_style(raw_text: str) -> Optional[str]:
    plain = _plain_text(raw_text)
    for pat, key in _CAPTION_PATTERNS:
        if pat.match(plain):
            return key
    return None


def _should_include_in_toc(text: str) -> bool:
    normalized = text.strip().lower()
    if normalized in {
        "table of contents",
        "abstract",
        "keywords",
        "key words",
        "acknowledgments",
        "acknowledgements",
        "references",
        "reference",
    }:
        return False
    return True


def _collect_static_toc_entries(blocks: List[Block]) -> List[Tuple[int, str]]:
    entries: List[Tuple[int, str]] = []
    for block in blocks:
        if block.kind != "heading":
            continue
        text = block.text.strip()
        effective_level = 2 if block.level == 1 and _NUMBERED_SECTION_HEADING_RE.match(text) else block.level
        if effective_level not in {2, 3, 4}:
            continue
        if not text or not _should_include_in_toc(text):
            continue
        entries.append((effective_level, text))
    return entries


def _insert_static_toc(doc: Document, entries: List[Tuple[int, str]]) -> None:
    for level, text in entries:
        p = doc.add_paragraph(style=_S["body"])
        if level == 3:
            p.paragraph_format.left_indent = Inches(0.32)
        elif level >= 4:
            p.paragraph_format.left_indent = Inches(0.58)
        apply_runs(p, parse_inline(text), spec_key="body")


# ---------------------------------------------------------------------------
# Document body clear
# ---------------------------------------------------------------------------

def _clear_body(doc: Document) -> None:
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _normalize_chart_heading(text: str) -> str:
    value = re.sub(r"^\s*\d+(?:\.\d+)*[.)]?\s*", "", text)
    value = re.sub(r"[^\w]+", " ", value.casefold(), flags=re.UNICODE)
    return re.sub(r"\s+", " ", value).strip()


_EXCLUDED_CHART_HEADING_KEYS = {
    "abstract", "keywords", "key words", "references", "reference list",
    "bibliography", "cited literature", "supporting information",
    "supplementary information", "table of contents",
}


def _expected_chart_headings(blocks: List[Block]) -> Dict[str, str]:
    expected: Dict[str, str] = {}
    for block_index, block in enumerate(blocks):
        if block.kind != "heading":
            continue
        numbered_h1 = block.level == 1 and _NUMBERED_SECTION_HEADING_RE.match(block.text.strip())
        effective_level = 2 if numbered_h1 else block.level
        if effective_level != 2:
            continue
        key = _normalize_chart_heading(block.text)
        if key and key not in _EXCLUDED_CHART_HEADING_KEYS:
            expected[key] = block.text.strip()
    return expected


def _manifest_image(base_dir: Path, entry: Any, label: str) -> Path:
    if not isinstance(entry, dict):
        raise ValueError(f"summary chart {label} manifest entry must be an object")
    rel = entry.get("path")
    expected_sha = entry.get("sha256")
    if not isinstance(rel, str) or not rel.strip():
        raise ValueError(f"summary chart {label} path is missing")
    rel_path = Path(rel)
    if rel_path.is_absolute():
        raise ValueError(f"summary chart {label} path must be relative")
    if rel_path.suffix.casefold() != ".png":
        raise ValueError(f"summary chart {label} image must use a .png path")
    if not isinstance(expected_sha, str) or not re.fullmatch(r"[0-9a-f]{64}", expected_sha):
        raise ValueError(f"summary chart {label} SHA-256 is invalid")
    image_path = (base_dir / rel).resolve()
    resolved_base = base_dir.resolve()
    if image_path != resolved_base and resolved_base not in image_path.parents:
        raise ValueError(f"summary chart {label} path escapes the draft directory")
    if not image_path.is_file():
        raise ValueError(f"summary chart {label} image is missing: {image_path}")
    if _sha256_file(image_path) != expected_sha:
        raise ValueError(f"summary chart {label} image hash does not match: {image_path}")
    try:
        with PILImage.open(image_path) as image:
            if image.format != "PNG":
                raise ValueError
            image.verify()
    except Exception:
        raise ValueError(
            f"summary chart {label} image is not a valid PNG: {image_path}"
        ) from None
    return image_path


def _load_summary_chart_bundle(md_path: Path, blocks: List[Block]) -> Optional[SummaryChartBundle]:
    manifest_path = md_path.parent / "review_summary_chart.json"
    if not manifest_path.exists():
        raise ValueError(
            f"summary chart manifest not found: {manifest_path}. "
            "The orchestrator's DOCX hard gate requires a current "
            "review_summary_chart.json beside the draft before export; run "
            "review-outline-summary-chart first."
        )
    try:
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise ValueError(f"cannot read summary chart manifest: {manifest_path}: {exc}") from exc
    stats = payload.get("stats") if isinstance(payload, dict) else None
    if not isinstance(stats, dict):
        raise ValueError("summary chart manifest is missing stats")
    if stats.get("generation_scope") not in {"full", "both"}:
        raise ValueError("overall review chart generation_scope must be 'full' or 'both' for DOCX export")
    draft_source = stats.get("draft_source")
    if not isinstance(draft_source, str) or not draft_source.strip():
        raise ValueError("summary chart draft_source is missing")
    source_path = Path(draft_source)
    if not source_path.is_absolute():
        source_path = md_path.parent / source_path
    if source_path.resolve() != md_path.resolve():
        raise ValueError("summary chart draft_source does not match the current Markdown draft")
    if stats.get("draft_sha256") != _sha256_file(md_path):
        raise ValueError("summary chart was not generated from the current Markdown draft")
    image_manifest = stats.get("image_manifest")
    if not isinstance(image_manifest, dict):
        raise ValueError("summary chart manifest is missing image_manifest")
    full = _manifest_image(md_path.parent, image_manifest.get("full"), "full")
    # The final release uses one global overview chart only.  Older manifests
    # may still include section images, but they are deliberately ignored so
    # no per-section mini-outline is inserted into the Word document.
    return SummaryChartBundle(full=full, sections={})


def _chart_width_inches(doc: Document, image_path: Path) -> float:
    usable_width = _usable_page_width_inches(doc)
    section = doc.sections[-1]
    usable_height = (
        section.page_height.inches
        - section.top_margin.inches
        - section.bottom_margin.inches
        - 0.75
    )
    with PILImage.open(image_path) as image:
        width_px, height_px = image.size
    if not width_px or not height_px:
        return usable_width
    return min(usable_width, usable_height * width_px / height_px)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def convert(md_path: Path, out_path: Path, template_path: Path) -> None:
    md_text, invalid_character_count = make_xml_compatible(
        md_path.read_text(encoding="utf-8")
    )
    if invalid_character_count:
        print(
            f"[md2docx] WARNING: replaced {invalid_character_count} unsupported "
            "control character(s) before XML generation."
        )
    md_text = normalize_mineru_latex(md_text)
    # Workflow comments carry internal paragraph/figure identifiers.  They are
    # useful to the dashboard before export, but are never manuscript prose.
    # Strip all HTML comments before tokenizing so metadata cannot leak into
    # the opening paragraph of the Word document.
    md_text = re.sub(r"<!--.*?-->", "", md_text, flags=re.S)
    blocks  = tokenize(md_text)
    toc_entries = _collect_static_toc_entries(blocks)
    # Only the global review overview is included.  Per-section mini-outline
    # images are intentionally excluded by _load_summary_chart_bundle().
    chart_bundle = _load_summary_chart_bundle(md_path, blocks)
    doc     = Document(str(template_path))
    _clear_body(doc)

    ctx: str           = "body"
    front_matter: bool = False
    inserted_toc_heading = False
    saw_toc_heading = False
    skipping_source_toc = False
    full_chart_inserted = False
    introduction_level: int | None = None
    inserted_section_charts: set[str] = set()
    chart_number = 0

    def insert_toc_once() -> None:
        nonlocal inserted_toc_heading
        if inserted_toc_heading:
            return
        _para(doc, "body", "h2", "Table of Contents", force_bold=True)
        _insert_static_toc(doc, toc_entries)
        inserted_toc_heading = True

    def insert_chart(image_path: Path, caption: str) -> None:
        nonlocal chart_number
        chart_number += 1
        paragraph = doc.add_paragraph(style=_S["body"])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(
            str(image_path),
            width=Inches(_chart_width_inches(doc, image_path)),
        )
        _para(doc, "chart", "chart", f"Chart {chart_number}. {caption}")

    for block_index, block in enumerate(blocks):

        if block.kind == "heading":
            plain_heading = block.text.strip().lower()
            numbered_h1_section = block.level == 1 and _NUMBERED_SECTION_HEADING_RE.match(block.text.strip())
            if plain_heading == "table of contents":
                saw_toc_heading = True
                skipping_source_toc = True
                insert_toc_once()
                continue
            elif block.level >= 2 and not inserted_toc_heading:
                insert_toc_once()
            skipping_source_toc = False
            effective_level = 2 if numbered_h1_section else block.level
            chart_heading_key = _normalize_chart_heading(block.text)
            if (
                chart_bundle is not None
                and not full_chart_inserted
                and should_insert_full_chart_before_heading(introduction_level, effective_level)
            ):
                insert_chart(chart_bundle.full, "Full-review structure summary.")
                full_chart_inserted = True
                introduction_level = None
            style_key, spec_key = _HEADING_FORMAT.get(effective_level, ("body", "body"))
            new_ctx = _section_ctx(block.text)
            ctx = new_ctx if new_ctx else "body"
            if block.level == 1 and not numbered_h1_section:
                front_matter = True
            elif effective_level >= 2:
                front_matter = False
            _para(doc, style_key, spec_key, block.text)
            if chart_bundle is not None and chart_heading_key in chart_bundle.sections:
                manifest_heading, image_path = chart_bundle.sections[chart_heading_key]
                insert_chart(image_path, f"Section structure summary: {manifest_heading}.")
                inserted_section_charts.add(chart_heading_key)
            if is_introduction_heading(block.text):
                introduction_level = effective_level

        elif block.kind == "paragraph":
            text  = block.text.strip()
            plain = _plain_text(text)

            # Bold-only section label  e.g. **Abstract**
            new_ctx = _section_ctx(plain)
            if new_ctx:
                skipping_source_toc = False
                if new_ctx in {"abstract", "keywords"} and not inserted_toc_heading and not saw_toc_heading:
                    insert_toc_once()
                ctx = new_ctx
                _para(doc, "body", "body", text, force_bold=True)
                continue
            if skipping_source_toc:
                continue

            # Front matter: author / affiliation
            if front_matter and ctx == "body":
                if _AFFIL_START.match(text):
                    _para(doc, "address", "address", text)
                else:
                    _para(doc, "author", "author", text)
                continue

            if ctx != "body":
                spec = ctx if ctx in _FONT_SPEC else "body"
                _para(doc, ctx, spec, text)
            else:
                cap = _caption_style(text)
                key = cap if cap else "body"
                _para(doc, key, key, text)

        elif block.kind == "indented_block":
            if skipping_source_toc:
                continue
            for raw_line in block.lines:
                text = raw_line.strip()
                if not text:
                    continue
                cap = _caption_style(text)
                key = cap if cap else ("references" if ctx == "references" else "body")
                spec = key if key in _FONT_SPEC else "body"
                _para(doc, key, spec, text)

        elif block.kind == "ref_def":
            if skipping_source_toc:
                continue
            _para(doc, "references", "references", block.text)

        elif block.kind == "list_item":
            if skipping_source_toc:
                continue
            indent = "  " * block.depth
            bullet = (f"{indent}- {block.text}"
                      if not block.ordered else f"{indent}{block.text}")
            if ctx == "references":
                _para(doc, "references", "references", bullet)
            else:
                _para(doc, "body", "body", bullet)

        elif block.kind == "code_block":
            if skipping_source_toc:
                continue
            p  = doc.add_paragraph(style=_S["body"])
            wr = p.add_run(block.code)
            wr.font.name = "Courier New"
            wr.font.size = Pt(9)

        elif block.kind == "math_block":
            if skipping_source_toc:
                continue
            p = doc.add_paragraph(style=_S["body"])
            _apply_math(p, block.latex)

        elif block.kind == "table":
            if skipping_source_toc:
                continue
            _add_table(doc, block.header, block.rows)

        elif block.kind == "image":
            if skipping_source_toc:
                continue
            img_path = Path(block.path)
            if not img_path.is_absolute():
                img_path = md_path.parent / img_path
            if img_path.exists():
                p = doc.add_paragraph(style=_S["body"])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(img_path), width=Inches(_usable_page_width_inches(doc)))
                next_block = blocks[block_index + 1] if block_index + 1 < len(blocks) else None
                next_caption = (
                    next_block is not None
                    and next_block.kind == "paragraph"
                    and _caption_style(_plain_text(next_block.text.strip())) is not None
                )
                # The drafting stage supplies a full caption immediately after
                # manuscript figures. Do not add a second generic caption.
                if block.alt and not next_caption:
                    kind = "Scheme" if block.alt.strip().lower().startswith("scheme") else "Figure"
                    _para(doc, "scheme" if kind == "Scheme" else "figure", "scheme" if kind == "Scheme" else "figure", f"{kind}. {block.alt}")
            else:
                continue

        elif block.kind == "hr":
            # Horizontal rules in review Markdown are section separators, not
            # desired visual borders in the final DOCX.
            continue

    if not inserted_toc_heading and not saw_toc_heading:
        insert_toc_once()

    if chart_bundle is not None:
        if not full_chart_inserted:
            raise ValueError("full-review summary chart could not be placed after Introduction")
        missing_sections = set(chart_bundle.sections) - inserted_section_charts
        if missing_sections:
            missing = ", ".join(
                chart_bundle.sections[key][0] for key in sorted(missing_sections)
            )
            raise ValueError(f"summary chart section headings were not found in Markdown: {missing}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[md2docx] Saved -> {out_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

_DEFAULT_TEMPLATE = Path(__file__).resolve().parent.parent / "review_template.docx"


def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="md2docx",
        description="Convert Markdown to DOCX using review_template.docx styles.",
    )
    p.add_argument("--input",    required=True, metavar="MD",   help="Input .md file")
    p.add_argument("--output",   required=True, metavar="DOCX", help="Output .docx file")
    p.add_argument("--template", default=str(_DEFAULT_TEMPLATE), metavar="DOCX",
                   help=f"Word template (default: {_DEFAULT_TEMPLATE})")
    return p


def main() -> None:
    args          = _build_parser().parse_args()
    md_path       = Path(args.input).resolve()
    out_path      = Path(args.output).resolve()
    template_path = Path(args.template).resolve()

    if not md_path.exists():
        raise SystemExit(f"[md2docx] ERROR: Input not found: {md_path}")
    if not template_path.exists():
        raise SystemExit(f"[md2docx] ERROR: Template not found: {template_path}")
    if not _LATEX_OK:
        print("[md2docx] INFO: latex2word not installed; "
              "using native Word text with subscript/superscript for supported MinerU formulas.")

    try:
        convert(md_path, out_path, template_path)
    except ValueError as exc:
        raise SystemExit(f"[md2docx] ERROR: {exc}") from None


if __name__ == "__main__":
    main()
