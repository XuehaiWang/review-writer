from __future__ import annotations

import base64
import tempfile
import unittest
import uuid
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from fastapi.testclient import TestClient
from sqlalchemy import create_engine, event, select
from sqlalchemy.orm import sessionmaker
from docx import Document

from review_writer_api.app import create_app
from review_writer_api.config import ApiSettings
from review_writer_api.database import Base, Project, User
from review_writer_api.domain_services.planning import (
    MATRIX_FACT_ENRICHMENT_CONTRACT_VERSION,
    TOPIC_PARTITION_BOUNDARY_LABEL,
    _matrix_classification_axes,
    _topic_outline_intent,
    _topic_partition_for_text,
    _topic_partition_for_row,
)
from review_writer_api.domain_services.library_index import EvidenceHit
from review_writer_api.security import Principal, Role
from review_writer_api.workflow_models import LibraryPaper


TEST_KEY = base64.urlsafe_b64encode(bytes(range(32))).decode("ascii").rstrip("=")


class PlanningV1Tests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        root = Path(self.temporary.name)
        database_url = f"sqlite+pysqlite:///{(root / 'planning.sqlite3').as_posix()}"
        self.engine = create_engine(database_url, connect_args={"check_same_thread": False})

        @event.listens_for(self.engine, "connect")
        def enable_foreign_keys(connection, _record):
            cursor = connection.cursor()
            cursor.execute("PRAGMA foreign_keys=ON")
            cursor.close()

        Base.metadata.create_all(self.engine)
        self.sessions = sessionmaker(bind=self.engine, expire_on_commit=False)
        with self.sessions.begin() as session:
            first = User(email="first@example.com", display_name="First", password_hash="hash")
            second = User(email="second@example.com", display_name="Second", password_hash="hash")
            session.add_all([first, second])
            session.flush()
            project = Project(user_id=first.id, slug="planning", topic="Copper allenation")
            hidden = Project(user_id=second.id, slug="hidden", topic="Hidden")
            session.add_all([project, hidden])
            session.flush()
            papers = []
            substrates = (
                "aromatic substrates",
                "small-molecule substrates",
                "biomolecular substrates",
            )
            catalysts = (
                "transition-metal catalysis",
                "organocatalysis",
                "photochemical methods",
            )
            reactions = (
                "cross-coupling",
                "addition reactions",
                "cyclization and annulation",
            )
            for index in range(1, 36):
                paper_id = f"P{index:03d}"
                structured_tags = {
                    "substrate": substrates[(index - 1) % len(substrates)],
                    "catalyst_or_method": catalysts[(index - 1) % len(catalysts)],
                    "reaction_type": reactions[(index - 1) % len(reactions)],
                }
                papers.append(
                    LibraryPaper(
                        user_id=first.id,
                        paper_id=paper_id,
                        content_sha256=f"{index:064x}",
                        original_filename=f"{paper_id}.pdf",
                        title=f"Paper {index}",
                        authors_json=[f"Author {index}"],
                        keywords_json=["allenation", "copper"],
                        tags_json=structured_tags,
                        metadata_json={
                            "paper_id": paper_id,
                            "title": {"value": f"Paper {index}"},
                            "authors": {"value": [f"Author {index}"]},
                            "keywords": {"value": ["allenation", "copper"]},
                            "abstract": {"value": f"Evidence for {paper_id}."},
                            "year": {"value": 2024},
                            "structured_tags": {
                                "value": structured_tags,
                                "human_checked": True,
                            },
                        },
                        pdf_relative_path=f"review-library/uploads/{paper_id}.pdf",
                        markdown_relative_path=f"review-library/markdown/{paper_id}.md",
                    )
                )
            session.add_all(papers)
            self.project_id = str(project.id)
            self.hidden_project_id = str(hidden.id)
            self.first = Principal(str(first.id), frozenset({Role.USER}), first.email)
            self.second = Principal(str(second.id), frozenset({Role.USER}), second.email)
        self.current = self.first
        settings = ApiSettings(
            review_root=root,
            deployment_mode="hosted",
            database_url=database_url,
            public_origin="http://testserver",
            credential_encryption_key=TEST_KEY,
            hosted_workspace_root=root / "users",
        )
        self.app = create_app(
            settings,
            principal_provider=lambda: self.current,
            session_factory_override=self.sessions,
        )
        self._seed_discovery(range(1, 36))

    def tearDown(self) -> None:
        Base.metadata.drop_all(self.engine)
        self.engine.dispose()
        self.temporary.cleanup()

    @staticmethod
    def headers() -> dict[str, str]:
        return {"Origin": "http://testserver"}

    def _review(self, selected: set[int]) -> dict:
        return {
            "project_id": self.project_id,
            "topic": "Copper allenation",
            "selection_mode": "explicit",
            "results": [
                {
                    "keyword": "allenation",
                    "keep": True,
                    "local_results": [
                        {
                            "paper_id": f"P{index:03d}",
                            "title": f"Paper {index}",
                            "score": 100 - index,
                            "role": "core_candidate",
                            "selected_for_matrix": index in selected,
                        }
                        for index in range(1, 36)
                    ],
                    "web_results": [],
                }
            ],
        }

    def _seed_discovery(self, selected) -> None:
        service = self.app.state.discovery_service
        repository = self.app.state.workflow_repository
        artifact, run = service._write_json_artifact(
            self.first,
            self.project_id,
            stage_id="discovery",
            logical_name="discovery/review.json",
            payload=self._review(set(selected)),
            make_current=False,
        )
        repository.save_discovery_atomically(
            self.first.user_id,
            self.project_id,
            artifact_id=artifact.id,
            run_id=run.id,
            expected_revision=0,
            status="review",
        )
        with TestClient(self.app) as client:
            response = client.post(
                f"/api/v1/projects/{self.project_id}/discovery/confirm",
                json={"revision": 1},
                headers=self.headers(),
            )
        self.assertEqual(200, response.status_code, response.text)

    def planning(self, client: TestClient) -> dict:
        response = client.get(f"/api/v1/projects/{self.project_id}/planning")
        self.assertEqual(200, response.status_code, response.text)
        return response.json()

    def choose_outline(self, client: TestClient, style: str = "substrate") -> dict:
        current = self.planning(client)
        response = client.put(
            f"/api/v1/projects/{self.project_id}/planning/outline",
            json={"revision": current["matrix_revision"], "outline_style": style},
            headers=self.headers(),
        )
        self.assertEqual(200, response.status_code, response.text)
        return response.json()

    def test_fact_contract_version_invalidates_old_cache_and_force_can_reextract(self) -> None:
        service = self.app.state.planning_service
        source = service.matrix_enrichment_payload(self.first, self.project_id)
        self.assertEqual(
            MATRIX_FACT_ENRICHMENT_CONTRACT_VERSION,
            source["fact_enrichment_contract_version"],
        )
        service.publish_matrix_enrichment(
            self.first,
            self.project_id,
            source,
            {
                "papers": [
                    {
                        "paper_id": paper["paper_id"],
                        "status": "complete",
                        "facts": [],
                        "failed_fields": [],
                    }
                    for paper in source["papers"]
                ]
            },
        )

        current = service.matrix_enrichment_payload(self.first, self.project_id)
        forced = service.matrix_enrichment_payload(
            self.first,
            self.project_id,
            force=True,
        )

        self.assertEqual(0, current["pending_paper_count"])
        self.assertEqual(source["paper_count"], forced["pending_paper_count"])
        self.assertTrue(forced["force_refresh"])

        with self.sessions.begin() as session:
            project = session.get(Project, uuid.UUID(self.project_id))
            project.model_tier = "luna"
        changed_model = service.matrix_enrichment_payload(
            self.first,
            self.project_id,
        )
        self.assertEqual(
            source["paper_count"], changed_model["pending_paper_count"]
        )
        self.assertEqual("gpt-5.6-luna", changed_model["actual_model_id"])

    def test_matrix_extraction_axes_ignore_runtime_coverage_fields(self) -> None:
        matrix = {
            "classification_axes": [
                {
                    "axis_id": "axis_01",
                    "label": "Method family",
                    "axis_role": "primary_organization",
                    "role_status": "evidence_confirmed",
                    "evidence_coverage": {"paper_count": 12, "partition_count": 3},
                    "partitions": [
                        {"partition_id": "method_a", "label": "Method A"}
                    ],
                }
            ]
        }

        axes = _matrix_classification_axes(matrix, [])

        self.assertNotIn("role_status", axes[0])
        self.assertNotIn("evidence_coverage", axes[0])
        self.assertIn("role_status", matrix["classification_axes"][0])

    def test_matrix_fact_retrieval_recovers_question_hits_inside_admitted_paper(self) -> None:
        class RecoveryIndex:
            enabled = True

            @staticmethod
            def summaries(_principal, paper_ids):
                return {
                    paper_id: {
                        "fulltext": "ready",
                        "chunker_version": "test-v1",
                        "source_lineage_hash": f"lineage-{paper_id}",
                    }
                    for paper_id in paper_ids
                    if paper_id == "P001"
                }

            @staticmethod
            def retrieve(
                _principal,
                _query,
                *,
                allowed_papers,
                term_groups=None,
                **_kwargs,
            ):
                groups = list(term_groups or [])
                if len(groups) != 1 or "substrate" not in groups[0]:
                    return []
                paper_id = allowed_papers[0]
                return [
                    EvidenceHit(
                        paper_id=paper_id,
                        chunk_id="results-001",
                        content="The substrate scope included aromatic examples.",
                        page_start=3,
                        page_end=3,
                        section_path=("Results",),
                        content_type="text",
                        asset_refs=(),
                        score=1.0,
                        match_reason="test",
                        is_neighbor=False,
                        index_id="index-001",
                        source_lineage_hash=f"lineage-{paper_id}",
                    )
                ]

        service = self.app.state.planning_service
        previous = service.library_index
        service.library_index = RecoveryIndex()
        try:
            payload = service.matrix_enrichment_payload(
                self.first, self.project_id, force=True
            )
        finally:
            service.library_index = previous

        paper = next(
            item for item in payload["papers"] if item["paper_id"] == "P001"
        )
        recovered = next(
            item
            for item in paper["evidence_candidates"]
            if item.get("chunk_id") == "results-001"
        )
        self.assertIn("object_input", recovered["question_ids"])
        self.assertIn(
            "admitted_paper_question_recovery", recovered["retrieval_passes"]
        )
        self.assertGreater(
            paper["retrieval_summary"]["relaxed_question_hit_count"], 0
        )

    def test_current_fact_refresh_is_an_idempotent_success_not_an_error(self) -> None:
        service = self.app.state.planning_service
        with patch.object(
            service,
            "matrix_enrichment_payload",
            return_value={
                "project_id": self.project_id,
                "pending_paper_count": 0,
                "fulltext_candidate_paper_count": 0,
            },
        ), TestClient(self.app) as client:
            response = client.post(
                f"/api/v1/projects/{self.project_id}/planning/matrix/enrichment/jobs",
                headers={**self.headers(), "Idempotency-Key": "current-facts"},
            )

        self.assertEqual(202, response.status_code, response.text)
        self.assertEqual("current", response.json()["status"])

    def test_all_fact_failures_require_explicit_limited_mode(self) -> None:
        service = self.app.state.planning_service
        current = service.get(self.first, self.project_id)
        rows = current["literature_matrix"]["rows"]
        source_payload = {
            "source_matrix_artifact_id": current["matrix_artifact_id"],
            "expected_matrix_revision": current["matrix_revision"],
            "papers": [
                {
                    "paper_id": row["paper_id"],
                    "source_fingerprint": f"failed-{row['paper_id']}",
                    "index_summary": {},
                    "evidence_candidates": [],
                }
                for row in rows
            ],
        }
        built = {
            "papers": [
                {
                    "paper_id": row["paper_id"],
                    "status": "failed",
                    "facts": [],
                    "failed_fields": ["all"],
                    "error": "provider unavailable",
                }
                for row in rows
            ]
        }
        service.publish_matrix_enrichment(
            self.first, self.project_id, source_payload, built
        )

        with TestClient(self.app) as client:
            blocked = self.planning(client)
            self.assertTrue(blocked["matrix_enrichment"]["planning_blocked"])
            self.choose_outline(client)
            blocked = self.planning(client)
            response = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint",
                json={"revision": blocked["blueprint_revision"]},
                headers=self.headers(),
            )
            self.assertEqual(409, response.status_code, response.text)

            limited = client.post(
                f"/api/v1/projects/{self.project_id}/planning/matrix/enrichment/limited-mode",
                json={"revision": blocked["matrix_revision"]},
                headers=self.headers(),
            )
            self.assertEqual(200, limited.status_code, limited.text)
            self.choose_outline(client)
            allowed = self.planning(client)
            response = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint",
                json={"revision": allowed["blueprint_revision"]},
                headers=self.headers(),
            )
            self.assertEqual(200, response.status_code, response.text)

    def test_fact_publish_tolerates_outline_revision_drift(self) -> None:
        service = self.app.state.planning_service
        current = service.get(self.first, self.project_id)
        rows = current["literature_matrix"]["rows"]
        source_payload = {
            "source_matrix_artifact_id": current["matrix_artifact_id"],
            "expected_matrix_revision": current["matrix_revision"],
            "papers": [
                {
                    "paper_id": row["paper_id"],
                    "source_fingerprint": f"fact-{row['paper_id']}",
                    "index_summary": {},
                    "evidence_candidates": [],
                }
                for row in rows
            ],
        }
        built = {
            "papers": [
                {
                    "paper_id": row["paper_id"],
                    "status": "failed",
                    "facts": [],
                    "failed_fields": ["all"],
                    "error": "no source-addressable evidence",
                }
                for row in rows
            ]
        }

        with TestClient(self.app) as client:
            self.choose_outline(client, "reaction")
            selected = self.planning(client)

        self.assertEqual(current["matrix_artifact_id"], selected["matrix_artifact_id"])
        self.assertGreater(selected["matrix_revision"], current["matrix_revision"])
        published = service.publish_matrix_enrichment(
            self.first, self.project_id, source_payload, built
        )
        reloaded = service.get(self.first, self.project_id)

        self.assertGreater(published["matrix_revision"], selected["matrix_revision"])
        self.assertEqual(35, reloaded["matrix_enrichment"]["counts"]["failed"])
        self.assertEqual(0, reloaded["matrix_enrichment"]["counts"]["pending"])
        self.assertTrue(reloaded["outline_current"])

    @staticmethod
    def isolated_reference_analysis(
        _principal,
        _project_id,
        *,
        candidate_id,
        safe_name,
        raw,
        matrix,
    ) -> dict:
        del safe_name, raw
        paper_ids = [row["paper_id"] for row in matrix["rows"]]
        representatives = paper_ids[:6]
        outline = (
            "# Selected Outline\n\n"
            "Scientific content source: current literature Matrix only.\n\n"
            "## Introduction and scope\n"
            f"Assigned papers: {', '.join(representatives)}.\n"
            "Purpose: define the current review scope.\n\n"
            "## 1. Copper allenation evidence\n"
            f"Assigned papers: {', '.join(paper_ids)}.\n"
            "Purpose: compare evidence from the current Matrix.\n\n"
            "## Conclusion and outlook\n"
            f"Assigned papers: {', '.join(representatives)}.\n"
            "Purpose: synthesize limitations and future directions.\n"
        )
        return {
            "candidate_id": candidate_id,
            "analysis_mode": "ai_style_only_transfer_v2",
            "content_source": "current_matrix_only",
            "reference_content_reused": False,
            "content_firewall": {
                "transfer_received_reference_text": False,
                "all_heading_levels_content_source": "current_matrix_only",
            },
            "reference_structure_metrics": {"heading_count": 3},
            "writing_style": {"organization_pattern": "progressive comparison"},
            "outline_md": outline,
        }

    def test_matrix_contains_entire_confirmed_selection(self) -> None:
        with TestClient(self.app) as client:
            payload = self.planning(client)
        self.assertEqual(35, len(payload["literature_matrix"]["rows"]))
        self.assertEqual(35, payload["matrix_sync"]["selected_paper_count"])
        self.assertNotIn("selection_fingerprint", payload["discovery_selection"])

    def test_reconfirmation_replaces_matrix_selection(self) -> None:
        with TestClient(self.app) as client:
            self.choose_outline(client)
            blueprint_revision = self.app.state.workflow_repository.get_stage_state(
                self.first.user_id, self.project_id, "blueprint"
            )
            if blueprint_revision is None:
                response = client.post(
                    f"/api/v1/projects/{self.project_id}/planning/blueprint",
                    json={"revision": 0},
                    headers=self.headers(),
                )
                self.assertEqual(200, response.status_code, response.text)
            previous_blueprint = (
                self.app.state.workflow_repository.get_current_artifact(
                    self.first.user_id,
                    self.project_id,
                    "blueprint/section_blueprint.json",
                )
            )
            review = client.get(f"/api/v1/projects/{self.project_id}/discovery").json()
            for row in review["results"][0]["local_results"]:
                row["selected_for_matrix"] = row["paper_id"] in {"P001", "P035"}
            saved = client.put(
                f"/api/v1/projects/{self.project_id}/discovery",
                json={"revision": review["revision"], "results": review["results"]},
                headers=self.headers(),
            ).json()
            confirmed = client.post(
                f"/api/v1/projects/{self.project_id}/discovery/confirm",
                json={"revision": saved["revision"]},
                headers=self.headers(),
            )
            self.assertEqual(200, confirmed.status_code, confirmed.text)
            planning = self.planning(client)
        self.assertEqual(["P001", "P035"], [row["paper_id"] for row in planning["literature_matrix"]["rows"]])
        self.assertEqual(
            previous_blueprint.id,
            self.app.state.workflow_repository.get_current_artifact(
                self.first.user_id, self.project_id, "blueprint/section_blueprint.json"
            ).id,
        )
        self.assertFalse(planning["outline_current"])
        self.assertFalse(planning["blueprint_current"])
        self.assertEqual(
            "stale",
            self.app.state.workflow_repository.get_stage_state(
                self.first.user_id, self.project_id, "blueprint"
            ).status,
        )

    def test_matrix_row_edit_uses_revision(self) -> None:
        with TestClient(self.app) as client:
            current = self.planning(client)
            response = client.put(
                f"/api/v1/projects/{self.project_id}/planning/matrix/P001",
                json={
                    "revision": current["matrix_revision"],
                    "main_content": "A" * 320,
                    "mark_complete": True,
                },
                headers=self.headers(),
            )
            self.assertEqual(200, response.status_code, response.text)
            stale = client.put(
                f"/api/v1/projects/{self.project_id}/planning/matrix/P002",
                json={"revision": current["matrix_revision"], "main_content": "changed"},
                headers=self.headers(),
            )
            reloaded = self.planning(client)
        self.assertEqual(409, stale.status_code, stale.text)
        rows = {row["paper_id"]: row for row in reloaded["literature_matrix"]["rows"]}
        self.assertEqual("full_reading_complete", rows["P001"]["matrix_status"])
        self.assertEqual("", rows["P002"]["main_content"])

    def test_builtin_outline_loads_editable_content(self) -> None:
        with TestClient(self.app) as client:
            selected = self.choose_outline(client, "reaction")
        self.assertIn("##", selected["selected_outline_md"])
        self.assertIn("## Introduction\nSection role: introduction", selected["selected_outline_md"])
        self.assertIn(
            "## Cross-category comparison and conclusion\nSection role: conclusion",
            selected["selected_outline_md"],
        )
        self.assertIn("## 1. Cross-coupling", selected["selected_outline_md"])
        self.assertIn("## 2. Addition reactions", selected["selected_outline_md"])
        self.assertIn("## 3. Cyclization and annulation", selected["selected_outline_md"])
        self.assertTrue(selected["outline_complete"])
        self.assertEqual("reaction", selected["outline_style"])

    def test_topic_outline_intent_uses_query_plan_and_explicit_partitions(self) -> None:
        topic = (
            "Please write a review on allenation-of-terminal-alkynes (ATA). "
            "Focus on mono-, 1,3-di-, and trisubstituted allenes. "
            "Organize the review by reaction type and catalytic/promoting system "
            "(Cu, Zn, Cd, Ti, etc.), and separately discuss racemic ATA and "
            "enantioselective ATA (EATA)."
        )
        discovery = {
            "query_plan": {
                "group_by": ["reaction_type", "catalyst_or_method"],
            }
        }

        intent = _topic_outline_intent(topic, discovery)

        self.assertTrue(intent["available"])
        self.assertEqual("reaction_type", intent["primary_axis"])
        self.assertEqual(["catalyst_or_method"], intent["secondary_axes"])
        self.assertEqual(
            ["racemic ATA", "enantioselective ATA (EATA)"],
            intent["partitions"],
        )
        self.assertEqual(intent["partitions"], intent["required_partitions"])
        self.assertEqual(["Cu", "Zn", "Cd", "Ti"], intent["named_systems"])
        self.assertEqual(intent["named_systems"], intent["comparison_dimensions"])
        self.assertEqual(
            [
                "mono-, 1,3-di-, and trisubstituted allenes",
            ],
            intent["focus_dimensions"],
        )
        self.assertEqual(intent["focus_dimensions"], intent["outcome_dimensions"])
        self.assertEqual(
            "source_bounded_model_or_section_contract",
            intent["partition_trace_policy"],
        )
        self.assertEqual(
            "reaction_type",
            intent["classification_contract"]["primary_axis_id"],
        )
        self.assertEqual(
            ["reaction_type", "catalyst_or_method"],
            [
                axis["axis_id"]
                for axis in intent["classification_contract"]["axes"][:2]
            ],
        )
        self.assertEqual(64, len(intent["classification_contract"]["fingerprint"]))

    def test_topic_intent_keeps_comparison_examples_out_of_required_partitions(self) -> None:
        topic = (
            "Organize the review by intervention type and compare age groups "
            "(children, adults, older adults), then separately discuss randomized "
            "evidence and observational evidence."
        )

        intent = _topic_outline_intent(topic, None)

        self.assertEqual(
            ["randomized evidence", "observational evidence"],
            intent["required_partitions"],
        )
        self.assertNotIn("children", intent["required_partitions"])
        self.assertNotIn("adults", intent["required_partitions"])

    def test_topic_intent_splits_repaired_stereochemistry_from_reaction_hierarchy(self) -> None:
        intent = _topic_outline_intent(
            (
                "Organize the review by reaction type and separately discuss "
                "racemic ATA and enantioselective ATA."
            ),
            {"query_plan": {"group_by": ["reaction_type"]}},
            [
                {
                    "axis_id": "stereochemical_regime",
                    "label": "Stereochemical regime",
                    "source_surface": (
                        "Organize the review by reaction type and separately "
                        "discuss racemic ATA and enantioselective ATA"
                    ),
                    "source_type": "explicit_topic",
                    "axis_role": "primary_organization",
                    "heading_requirement": "primary_heading",
                    "semantic_repair": {"status": "auto_repaired"},
                    "partitions": [
                        {"partition_id": "racemic", "label": "racemic ATA"},
                        {
                            "partition_id": "enantioselective",
                            "label": "enantioselective ATA",
                        },
                    ],
                }
            ],
        )

        self.assertEqual("reaction_type", intent["primary_axis"])
        self.assertEqual(
            ["stereochemical_regime"], intent["secondary_axes"]
        )
        self.assertEqual(
            ["racemic ATA", "enantioselective ATA"],
            intent["required_partitions"],
        )
        self.assertEqual(
            "required_independent_discussion",
            intent["classification_axes"][1]["axis_role"],
        )
        self.assertEqual(
            intent["classification_contract"]["axes"],
            intent["classification_axes"],
        )

    def test_topic_intent_reads_examples_for_other_chemistry_axes(self) -> None:
        topic = (
            "Review the transformation categorized by substrates "
            "(aryl halides, alkenes, organoboron reagents, etc.)."
        )
        discovery = {"query_plan": {"group_by": ["substrate"]}}

        intent = _topic_outline_intent(topic, discovery)

        self.assertEqual("substrate", intent["primary_axis"])
        self.assertEqual(
            ["aryl halides", "alkenes", "organoboron reagents"],
            intent["axis_examples"]["substrate"],
        )
        self.assertEqual([], intent["required_partitions"])

    def test_topic_guided_outline_supports_product_as_primary_axis(self) -> None:
        service = self.app.state.planning_service
        rows = [{"paper_id": "P001"}, {"paper_id": "P002"}]
        intent = {
            "available": True,
            "primary_axis": "product",
            "secondary_axes": ["catalyst_or_method"],
            "required_partitions": [],
        }

        outline = service._topic_outline_document(
            rows,
            tags_by_paper={
                "P001": {"product": "heterocycles"},
                "P002": {"product": "pharmaceutical compounds"},
            },
            text_by_paper={
                "P001": "Organocatalysis furnished a heterocyclic compound.",
                "P002": "Enzymatic methods furnished a pharmaceutical product.",
            },
            taxonomy_profile="chemistry_general",
            intent=intent,
        )

        self.assertIn("Primary structure: Topic-guided (product class", outline)
        self.assertIn("Heterocycles", outline)
        self.assertIn("Pharmaceutical compounds", outline)

    def test_topic_partition_matching_uses_declared_terms_without_domain_defaults(self) -> None:
        partitions = ["photochemical conditions", "electrochemical conditions"]

        self.assertEqual(
            "photochemical conditions",
            _topic_partition_for_text(
                "The reaction was performed under photochemical conditions.",
                partitions,
            ),
        )
        self.assertEqual(
            "",
            _topic_partition_for_text(
                "The source reports thermal activation but no requested partition.",
                partitions,
            ),
        )

    def test_topic_guided_outline_combines_partitions_and_matrix_axis(self) -> None:
        service = self.app.state.planning_service
        rows = [{"paper_id": f"P{index:03d}"} for index in range(1, 5)]
        tags_by_paper = {
            "P001": {"reaction_type": "three-component coupling"},
            "P002": {"reaction_type": "three-component coupling"},
            "P003": {"reaction_type": "homologation"},
            "P004": {"reaction_type": "homologation"},
        }
        text_by_paper = {
            "P001": "Racemic Cu-promoted terminal alkyne allenation afforded an allene.",
            "P002": "Enantioselective Cu-catalyzed ATA afforded 95% ee.",
            "P003": "Racemic zinc-promoted homologation of a terminal alkyne.",
            "P004": "Asymmetric homologation gave an enantioenriched allene.",
        }
        intent = {
            "available": True,
            "primary_axis": "reaction_type",
            "secondary_axes": ["catalyst_or_method"],
            "required_partitions": ["racemic ATA", "enantioselective ATA (EATA)"],
            "comparison_dimensions": ["Cu", "Zn", "Cd", "Ti"],
            "focus_dimensions": [
                "monosubstituted allenes",
                "1,3-disubstituted allenes",
                "trisubstituted allenes",
            ],
        }

        outline = service._topic_outline_document(
            rows,
            tags_by_paper=tags_by_paper,
            text_by_paper=text_by_paper,
            taxonomy_profile="chemistry_general",
            intent=intent,
        )

        self.assertIn(
            "Racemic ATA — Three-component coupling",
            outline,
        )
        self.assertIn(
            "Enantioselective ATA (EATA) — Three-component coupling",
            outline,
        )
        self.assertIn("Assigned papers: P001.", outline)
        self.assertIn("Assigned papers: P002.", outline)
        self.assertIn("catalytic or promoting system", outline)
        self.assertIn(
            "Focus dimensions: monosubstituted allenes, "
            "1,3-disubstituted allenes, trisubstituted allenes.",
            outline,
        )

    def test_topic_guided_outline_prefers_evidence_bound_model_partition(self) -> None:
        service = self.app.state.planning_service
        rows = [
            {
                "paper_id": "P001",
                "topic_partition_classification": {
                    "status": "classified",
                    "partition": "randomized evidence",
                    "confidence": 0.92,
                    "evidence_refs": [{"evidence_key": "sha256:source"}],
                },
            }
        ]
        outline = service._topic_outline_document(
            rows,
            tags_by_paper={"P001": {"reaction_type": "controlled comparison"}},
            text_by_paper={
                "P001": "The title and abstract use no literal partition label."
            },
            taxonomy_profile="general",
            intent={
                "available": True,
                "primary_axis": "reaction_type",
                "secondary_axes": [],
                "required_partitions": [
                    "randomized evidence",
                    "observational evidence",
                ],
            },
        )

        self.assertIn(
            "Randomized evidence — Controlled comparison",
            outline,
        )
        self.assertIn("Topic partition: randomized evidence.", outline)

    def test_completed_model_boundary_is_not_overruled_by_keyword_match(self) -> None:
        routed = _topic_partition_for_row(
            {
                "topic_partition_classification": {
                    "status": "boundary",
                    "partition": "",
                    "confidence": 0.42,
                    "evidence_refs": [],
                }
            },
            ["randomized evidence", "observational evidence"],
            "The related-work paragraph mentions randomized evidence.",
        )

        self.assertEqual("", routed)

    def test_formal_matrix_tag_routes_before_legacy_topic_text_match(self) -> None:
        routed = _topic_partition_for_row(
            {
                "evidence_backed_tags": {
                    "study_design": [
                        {
                            "partition_label": "randomized evidence",
                            "confidence": 0.93,
                            "fact_ids": ["MF-1"],
                            "evidence_refs": [{"evidence_key": "sha256:formal"}],
                        }
                    ]
                }
            },
            ["randomized evidence", "observational evidence"],
            "The background mentions observational evidence.",
        )
        self.assertEqual("randomized evidence", routed)

    def test_topic_guided_outline_uses_configured_taxonomy_without_topic_branch(self) -> None:
        service = self.app.state.planning_service
        rows = [{"paper_id": f"P{index:03d}"} for index in range(1, 4)]
        text_by_paper = {
            "P001": "Transition metal catalysis enabled a cross-coupling reaction.",
            "P002": "Visible light photochemical conditions enabled a cycloaddition.",
            "P003": "An electrochemical oxidation furnished the target product.",
        }
        intent = {
            "available": True,
            "primary_axis": "reaction_type",
            "secondary_axes": ["catalyst_or_method"],
            "partitions": [],
            "comparison_dimensions": ["operating conditions"],
        }

        outline = service._topic_outline_document(
            rows,
            tags_by_paper={},
            text_by_paper=text_by_paper,
            taxonomy_profile="chemistry_general",
            intent=intent,
        )

        self.assertIn("Cross-coupling", outline)
        self.assertIn("Cyclization and annulation", outline)
        self.assertIn("Oxidation and reduction", outline)
        self.assertIn("transition-metal catalysis", outline)
        self.assertIn("photochemical methods", outline)
        self.assertIn("electrochemical methods", outline)

    def test_reselecting_current_outline_is_idempotent(self) -> None:
        with TestClient(self.app) as client:
            selected = self.choose_outline(client, "reaction")
            repeated = client.put(
                f"/api/v1/projects/{self.project_id}/planning/outline",
                json={
                    "revision": selected["matrix_revision"],
                    "outline_style": "reaction",
                },
                headers=self.headers(),
            )
        self.assertEqual(200, repeated.status_code, repeated.text)
        payload = repeated.json()
        self.assertTrue(payload["unchanged"])
        self.assertEqual(selected["matrix_revision"], payload["matrix_revision"])
        self.assertEqual(selected["outline_artifact_id"], payload["outline_artifact_id"])

    def test_builtin_outline_styles_use_distinct_metadata_axes(self) -> None:
        with TestClient(self.app) as client:
            substrate = self.choose_outline(client, "substrate")["selected_outline_md"]
            catalyst = self.choose_outline(client, "catalyst")["selected_outline_md"]
            reaction = self.choose_outline(client, "reaction")["selected_outline_md"]
        self.assertIn("## 1. Aromatic substrates", substrate)
        self.assertIn("## 1. Transition-metal catalysis", catalyst)
        self.assertIn("## 1. Cross-coupling", reaction)
        self.assertNotEqual(substrate, catalyst)
        self.assertNotEqual(catalyst, reaction)

    def test_chemistry_substrate_outline_can_group_common_allene_precursors(self) -> None:
        service = self.app.state.planning_service
        rows = [{"paper_id": f"P{index:03d}"} for index in range(1, 8)]
        text_by_paper = {
            "P001": "enantioselective conversion of a propargylic alcohol substrate",
            "P002": "scope of substituted propargyl alcohol starting materials",
            "P003": "catalytic synthesis from a terminal alkyne",
            "P004": "three-component reaction using 1-alkynes",
            "P005": "cyclization of a conjugated enyne",
            "P006": "asymmetric transformation of substituted 1,3-enynes",
            "P007": "terminal alkynes as allene precursors; the abstract later compares propargylic alcohol chemistry",
        }

        groups = service._semantic_outline_groups(
            rows,
            text_by_paper,
            tag_key="substrate",
            taxonomy_profile="chemistry_general",
        )

        self.assertEqual(["P001", "P002"], groups["propargylic alcohols"])
        self.assertEqual(["P003", "P004", "P007"], groups["terminal alkynes"])
        self.assertEqual(["P005", "P006"], groups["conjugated enynes"])

    def test_chemistry_outline_reroutes_allene_evidence_without_catch_all(self) -> None:
        service = self.app.state.planning_service
        rows = [{"paper_id": f"P{index:03d}"} for index in range(1, 7)]
        text_by_paper = {
            "P001": "enantioselective isomerization of 3-alkynoates to chiral allenoates",
            "P002": "this review will highlight allenes in catalytic asymmetric synthesis",
            "P003": "phase-transfer functionalization of 1-alkylallene-1,3-dicarboxylates",
            "P004": "palladium synthesis of axially chiral (allenylmethyl)silanes",
            "P005": "hydroboration of but-1-en-3-ynes to axially chiral allenylboranes",
            "P006": "resolution of an allene hydrocarbon into optical antipodes",
        }

        outline = service._outline_document(
            "substrate",
            rows,
            tags_by_paper={paper_id: {} for paper_id in text_by_paper},
            text_by_paper=text_by_paper,
            taxonomy_profile="chemistry_general",
        )

        self.assertNotIn("Routing required", outline)
        self.assertIn("Context papers: P002.", outline)
        self.assertIn("## 1. Alkynoates", outline)
        self.assertIn("Preformed substituted allenes", outline)
        self.assertIn("Silyl-substituted diene precursors", outline)
        self.assertIn("Enynes", outline)

    def test_chemistry_outline_routes_resolution_and_enynamide_titles(self) -> None:
        service = self.app.state.planning_service
        rows = [{"paper_id": "P001"}, {"paper_id": "P002"}]
        groups = service._semantic_outline_groups(
            rows,
            {
                "P001": "chemoenzymatic dynamic kinetic resolution of axially chiral allenes",
                "P002": "rhodium-catalyzed 1,6-addition of arylboronic acids to enynamides",
            },
            tag_key="substrate",
            taxonomy_profile="chemistry_general",
        )

        self.assertEqual(["P001"], groups["preformed substituted allenes"])
        self.assertEqual(["P002"], groups["enynes"])
        self.assertNotIn("Routing required — reassign these papers", groups)

    def test_generated_routing_placeholder_is_repaired_before_blueprint(self) -> None:
        service = self.app.state.planning_service
        sections = [
            {
                "title": "Introduction",
                "section_role": "introduction",
                "paper_ids": [],
            },
            {
                "title": "preformed substituted allenes",
                "section_role": "body",
                "paper_ids": ["P003"],
            },
            {
                "title": "enynes",
                "section_role": "body",
                "paper_ids": ["P004"],
            },
            {
                "title": "Routing required — reassign these papers",
                "section_role": "body",
                "paper_ids": ["P001", "P002"],
            },
            {
                "title": "Conclusion",
                "section_role": "conclusion",
                "paper_ids": [],
            },
        ]
        repaired, adjustments = service._auto_repair_generated_routing_sections(
            sections,
            [{"paper_id": f"P00{index}"} for index in range(1, 5)],
            {
                "P001": "chemoenzymatic dynamic kinetic resolution of axially chiral allenes",
                "P002": "enantioselective 1,6-addition to enynamides",
            },
            outline_style="substrate",
            taxonomy_profile="chemistry_general",
        )

        by_title = {section["title"]: section for section in repaired}
        self.assertNotIn("Routing required — reassign these papers", by_title)
        self.assertEqual(
            ["P003", "P001"],
            by_title["preformed substituted allenes"]["paper_ids"],
        )
        self.assertEqual(["P004", "P002"], by_title["enynes"]["paper_ids"])
        self.assertEqual(2, len(adjustments))

    def test_generated_boundary_section_is_repaired_again_from_scientific_objects(self) -> None:
        service = self.app.state.planning_service
        repaired, adjustments = service._auto_repair_generated_routing_sections(
            [
                {"title": "Introduction", "section_role": "introduction", "paper_ids": []},
                {
                    "title": "Cross-category evidence and boundary cases",
                    "section_role": "body",
                    "paper_ids": ["P001", "P002"],
                },
                {"title": "Conclusion", "section_role": "conclusion", "paper_ids": []},
            ],
            [{"paper_id": "P001"}, {"paper_id": "P002"}],
            {
                "P001": "propargylic benzoates are converted to axially chiral allenes",
                "P002": "activated enynes are converted to axially chiral allenes",
            },
            outline_style="substrate",
            taxonomy_profile="chemistry_general",
        )

        by_title = {section["title"]: section for section in repaired}
        self.assertNotIn("Cross-category evidence and boundary cases", by_title)
        self.assertEqual(["P001"], by_title["activated propargylic derivatives"]["paper_ids"])
        self.assertEqual(["P002"], by_title["enynes"]["paper_ids"])
        self.assertEqual(2, len(adjustments))

    def test_unresolved_primary_study_is_not_relabelled_as_introduction_context(self) -> None:
        service = self.app.state.planning_service
        repaired, adjustments = service._auto_repair_generated_routing_sections(
            [
                {
                    "title": "Introduction",
                    "section_role": "introduction",
                    "paper_ids": [],
                    "context_paper_ids": [],
                },
                {
                    "title": "Routing required — reassign these papers",
                    "section_role": "body",
                    "paper_ids": ["P001"],
                },
                {
                    "title": "Conclusion",
                    "section_role": "conclusion",
                    "paper_ids": [],
                },
            ],
            [{"paper_id": "P001"}],
            {"P001": "A selected primary experiment with no supported route."},
            outline_style="reaction",
            taxonomy_profile="general_academic",
        )

        introduction = next(
            section
            for section in repaired
            if section["section_role"] == "introduction"
        )
        self.assertEqual([], introduction["context_paper_ids"])
        self.assertEqual("unresolved_classification_retained", adjustments[0]["method"])
        self.assertEqual(["P001"], adjustments[0]["paper_ids"])

    def test_topic_guided_repair_preserves_non_default_primary_axis(self) -> None:
        service = self.app.state.planning_service
        repaired, _adjustments = service._auto_repair_generated_routing_sections(
            [
                {
                    "title": "Routing required — reassign these papers",
                    "section_role": "body",
                    "paper_ids": ["P001"],
                }
            ],
            [{"paper_id": "P001"}],
            {"P001": "The reaction furnished a heterocyclic compound."},
            outline_style="topic-guided",
            taxonomy_profile="chemistry_general",
            tag_key_override="product",
            axis_label_override="product class",
        )

        self.assertEqual("heterocycles", repaired[0]["title"])
        self.assertIn("product class", repaired[0]["purpose"])

    def test_account_language_is_contextual_evidence(self) -> None:
        contextual = self.app.state.planning_service._contextual_outline_paper_ids(
            [{"paper_id": "P001"}, {"paper_id": "P002"}],
            {"P001": {}, "P002": {}},
            {
                "P001": "The account concerns palladium-catalyzed cyclization reactions.",
                "P002": "We report a controlled primary study.",
            },
        )

        self.assertEqual(["P001"], contextual)

    def test_generated_body_is_realigned_from_input_object_not_product_title(self) -> None:
        repaired, adjustments = (
            self.app.state.planning_service._realign_generated_body_sections(
                [
                    {"title": "Introduction", "section_role": "introduction", "paper_ids": []},
                    {
                        "title": "allenoates",
                        "section_role": "body",
                        "paper_ids": ["P001"],
                    },
                    {"title": "Conclusion", "section_role": "conclusion", "paper_ids": []},
                ],
                [
                    {
                        "paper_id": "P001",
                        "scientific_facts": [
                            {
                                "field_id": "object_input",
                                "value": "nitroalkanes and activated enynes",
                            }
                        ],
                    }
                ],
                {
                    "P001": (
                        "nitroalkanes and activated enynes are the input objects. "
                        "enantioselective synthesis of axially chiral allenes and allenoates"
                    )
                },
                outline_style="substrate",
                taxonomy_profile="chemistry_general",
            )
        )

        by_title = {section["title"]: section for section in repaired}
        self.assertNotIn("allenoates", by_title)
        self.assertEqual(["P001"], by_title["enynes"]["paper_ids"])
        self.assertEqual("scientific_object_reassignment", adjustments[0]["method"])

    def test_topic_partition_survives_scientific_object_realignment(self) -> None:
        repaired, _adjustments = (
            self.app.state.planning_service._realign_generated_body_sections(
                [
                    {
                        "title": "Randomized evidence — Allenoates",
                        "section_role": "body",
                        "topic_partition": "randomized evidence",
                        "paper_ids": ["P001"],
                    }
                ],
                [
                    {
                        "paper_id": "P001",
                        "scientific_facts": [
                            {
                                "field_id": "object_input",
                                "value": "activated enynes",
                            }
                        ],
                    }
                ],
                {
                    "P001": (
                        "activated enynes furnished an axially chiral allene product"
                    )
                },
                outline_style="topic-guided",
                taxonomy_profile="chemistry_general",
                tag_key_override="substrate",
                axis_label_override="substrate class",
            )
        )

        self.assertEqual(1, len(repaired))
        self.assertEqual("randomized evidence", repaired[0]["topic_partition"])
        self.assertEqual(
            "Randomized evidence — Enynes",
            repaired[0]["title"],
        )

    def test_topic_boundary_rationale_survives_primary_axis_realignment(self) -> None:
        rationale = (
            "The source does not positively establish one requested Topic partition."
        )
        repaired, _adjustments = (
            self.app.state.planning_service._realign_generated_body_sections(
                [
                    {
                        "title": "Topic-partition boundary cases — ATA",
                        "section_role": "body",
                        "topic_partition": TOPIC_PARTITION_BOUNDARY_LABEL,
                        "boundary_rationale": rationale,
                        "paper_ids": ["P001"],
                    }
                ],
                [
                    {
                        "paper_id": "P001",
                        "scientific_facts": [
                            {
                                "field_id": "transformation",
                                "value": "terminal alkyne allenation",
                            }
                        ],
                    }
                ],
                {"P001": "terminal alkyne allenation furnished an allene"},
                outline_style="topic-guided",
                taxonomy_profile="chemistry_general",
                tag_key_override="reaction_type",
                axis_label_override="reaction type",
            )
        )

        self.assertEqual(1, len(repaired))
        self.assertEqual(
            TOPIC_PARTITION_BOUNDARY_LABEL,
            repaired[0]["topic_partition"],
        )
        self.assertEqual("Topic-partition boundary cases — ATA", repaired[0]["title"])
        self.assertEqual(rationale, repaired[0]["boundary_rationale"])

    def test_outline_sources_prefer_confirmed_and_ignore_legacy_automatic_tags(self) -> None:
        service = self.app.state.planning_service
        confirmed_tags, _ = service._outline_sources(
            self.first,
            [
                {
                    "paper_id": "P001",
                    "project_tag_review_status": "confirmed",
                    "project_tags": {
                        "reaction_type": ["project-specific transformation"]
                    },
                }
            ],
        )
        pending_tags, _ = service._outline_sources(
            self.first,
            [
                {
                    "paper_id": "P001",
                    "project_tag_review_status": "pending",
                    "project_tags": {"reaction_type": ["unreviewed suggestion"]},
                }
            ],
        )
        automatic_tags, _ = service._outline_sources(
            self.first,
            [
                {
                    "paper_id": "P001",
                    "project_tag_review_status": "automatic",
                    "project_tags": {
                        "reaction_type": ["automatically assessed transformation"]
                    },
                }
            ],
        )
        self.assertEqual(
            ["project-specific transformation"],
            confirmed_tags["P001"]["reaction_type"],
        )
        self.assertEqual("cross-coupling", pending_tags["P001"]["reaction_type"])
        self.assertEqual("cross-coupling", automatic_tags["P001"]["reaction_type"])

    def test_outline_sources_use_evidence_bounded_agent_routing_as_fallback(self) -> None:
        with patch(
            "review_writer_api.domain_services.planning.verified_structured_tags",
            return_value={},
        ):
            tags, _ = self.app.state.planning_service._outline_sources(
                self.first,
                [
                    {
                        "paper_id": "P001",
                        "routing_recommendation": {
                            "axis_id": "reaction_type",
                            "status": "classified",
                            "label": "aldehyde-based three-component ATA",
                            "confidence": 0.96,
                            "evidence_refs": [{"evidence_key": "sha256:route"}],
                        },
                    }
                ],
            )

        self.assertEqual(
            ["aldehyde-based three-component ATA"],
            tags["P001"]["reaction_type"],
        )

    def test_outline_sources_ignore_unverified_library_tags(self) -> None:
        with self.sessions.begin() as session:
            paper = session.scalar(
                select(LibraryPaper).where(LibraryPaper.paper_id == "P001")
            )
            metadata = dict(paper.metadata_json)
            structured = dict(metadata["structured_tags"])
            structured["human_checked"] = False
            metadata["structured_tags"] = structured
            paper.metadata_json = metadata

        tags, _ = self.app.state.planning_service._outline_sources(
            self.first,
            [{"paper_id": "P001", "project_tag_review_status": "pending"}],
        )

        self.assertEqual({}, tags["P001"])

    def test_custom_outline_starts_blank(self) -> None:
        with TestClient(self.app) as client:
            selected = self.choose_outline(client, "custom")
        self.assertEqual("", selected["selected_outline_md"])
        self.assertFalse(selected["outline_complete"])

    def test_manual_outline_save_versions_content(self) -> None:
        with TestClient(self.app) as client:
            selected = self.choose_outline(client, "custom")
            outline = "# Review\n\n## 1. Introduction\nAssigned papers: P001, P002.\nPurpose: scope.\n"
            response = client.put(
                f"/api/v1/projects/{self.project_id}/planning/outline",
                json={
                    "revision": selected["matrix_revision"],
                    "outline_style": "custom",
                    "outline_md": outline,
                },
                headers=self.headers(),
            )
            self.assertEqual(200, response.status_code, response.text)
            saved = response.json()
            stale = client.put(
                f"/api/v1/projects/{self.project_id}/planning/outline",
                json={
                    "revision": selected["matrix_revision"],
                    "outline_style": "custom",
                    "outline_md": outline.replace("scope", "stale"),
                },
                headers=self.headers(),
            )
            reloaded = self.planning(client)
        self.assertEqual(409, stale.status_code)
        self.assertEqual(saved["outline_artifact_id"], reloaded["outline_selection"]["artifact_id"])
        self.assertIn("Purpose: scope.", reloaded["selected_outline_md"])

    def test_saved_outline_appears_in_comparison(self) -> None:
        with TestClient(self.app) as client:
            self.choose_outline(client, "catalyst")
            payload = self.planning(client)
        candidates = {item["candidate_id"]: item for item in payload["outline_candidates"]}
        self.assertIn("saved-current", candidates)
        self.assertEqual(payload["selected_outline_md"], candidates["saved-current"]["outline_md"])

    def test_reference_outline_is_registered(self) -> None:
        raw = "# Reference\n\n## 1. Mechanisms\nAssigned papers: P001.\nPurpose: compare.\n".encode()
        with patch.object(
            self.app.state.planning_service,
            "_analyze_reference_document",
            side_effect=self.isolated_reference_analysis,
        ):
            with TestClient(self.app) as client:
                current = self.planning(client)
                response = client.post(
                    f"/api/v1/projects/{self.project_id}/planning/reference-outlines",
                    json={
                        "revision": current["matrix_revision"],
                        "filename": "reference.md",
                        "content_base64": base64.b64encode(raw).decode(),
                    },
                    headers=self.headers(),
                )
                self.assertEqual(201, response.status_code, response.text)
                payload = self.planning(client)
                candidate_id = response.json()["candidate"]["candidate_id"]
                selected_response = client.put(
                    f"/api/v1/projects/{self.project_id}/planning/outline",
                    json={
                        "revision": payload["matrix_revision"],
                        "outline_style": f"reference:{candidate_id}",
                    },
                    headers=self.headers(),
                )
                self.assertEqual(200, selected_response.status_code, selected_response.text)
        candidate = response.json()["candidate"]
        self.assertTrue(candidate["source_artifact_id"])
        self.assertEqual("current_matrix_only", candidate["content_source"])
        self.assertFalse(candidate["reference_content_reused"])
        self.assertNotIn("Mechanisms", candidate["outline_md"])
        self.assertIn(candidate["candidate_id"], {item["candidate_id"] for item in payload["reference_outline_candidates"]})
        self.assertNotIn("Mechanisms", selected_response.json()["selected_outline_md"])

    def test_legacy_reference_candidate_fails_content_isolation(self) -> None:
        service = self.app.state.planning_service
        self.assertFalse(
            service._reference_candidate_is_isolated(
                {
                    "analysis_mode": "heading_extraction",
                    "outline_md": "## Source heading",
                }
            )
        )

    def test_reference_docx_content_is_not_used_as_candidate_headings(self) -> None:
        document = Document()
        document.add_heading("1. Mechanistic organization", level=1)
        document.add_paragraph("Reference discussion.")
        stream = BytesIO()
        document.save(stream)
        with patch.object(
            self.app.state.planning_service,
            "_analyze_reference_document",
            side_effect=self.isolated_reference_analysis,
        ):
            with TestClient(self.app) as client:
                current = self.planning(client)
                response = client.post(
                    f"/api/v1/projects/{self.project_id}/planning/reference-outlines",
                    json={
                        "revision": current["matrix_revision"],
                        "filename": "reference.docx",
                        "content_base64": base64.b64encode(stream.getvalue()).decode(),
                    },
                    headers=self.headers(),
                )
        self.assertEqual(201, response.status_code, response.text)
        candidate = response.json()["candidate"]
        self.assertEqual("ai_style_only_transfer_v2", candidate["analysis_mode"])
        self.assertNotIn("Mechanistic organization", candidate["outline_md"])
        self.assertIn("Copper allenation evidence", candidate["outline_md"])

    def test_blueprint_uses_current_matrix_and_outline(self) -> None:
        with TestClient(self.app) as client:
            selected = self.choose_outline(client, "substrate")
            response = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint",
                json={"revision": 0},
                headers=self.headers(),
            )
            self.assertEqual(200, response.status_code, response.text)
            blueprint = response.json()["section_blueprint"]
        matrix_ids = {f"P{index:03d}" for index in range(1, 36)}
        assigned = {paper_id for section in blueprint["sections"] for paper_id in section["major_papers"]}
        self.assertTrue(assigned)
        self.assertLessEqual(assigned, matrix_ids)
        primary_occurrences = [
            paper_id
            for section in blueprint["sections"]
            for paper_id in section["primary_papers"]
        ]
        self.assertEqual(len(primary_occurrences), len(set(primary_occurrences)))
        introduction = next(
            section
            for section in blueprint["sections"]
            if section["section_role"] == "introduction"
        )
        conclusion = next(
            section
            for section in blueprint["sections"]
            if section["section_role"] == "conclusion"
        )
        self.assertEqual([], introduction["major_papers"])
        self.assertEqual([], conclusion["major_papers"])
        self.assertTrue(introduction["supporting_papers"])
        self.assertTrue(conclusion["supporting_papers"])
        self.assertTrue(
            blueprint["paper_assignment_policy"][
                "introduction_and_conclusion_are_synthesis_only"
            ]
        )
        self.assertEqual(selected["outline_artifact_id"], blueprint["source_outline_artifact_id"])
        self.assertEqual(
            blueprint["classification_contract"]["fingerprint"],
            blueprint["classification_basis"]["axis_contract_fingerprint"],
        )
        self.assertEqual(
            blueprint["classification_contract"]["fingerprint"],
            blueprint["classification_contract_lineage"]["effective_fingerprint"],
        )

    def test_planning_bundle_exposes_scope_and_synthesis_requirements(self) -> None:
        with TestClient(self.app) as client:
            self.choose_outline(client, "reaction")
            generated = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint",
                json={"revision": 0},
                headers=self.headers(),
            )
        self.assertEqual(200, generated.status_code, generated.text)
        blueprint = generated.json()["section_blueprint"]
        self.assertTrue(blueprint["scope_diagnostics"]["can_confirm"])
        self.assertTrue(blueprint["taxonomy_diagnostics"]["can_confirm"])
        self.assertEqual(
            "reaction_strategy",
            blueprint["scope_contract"]["primary_navigation_axis"],
        )
        self.assertTrue(blueprint["synthesis_requirements"])
        self.assertTrue(
            all("academic_contract" in section for section in blueprint["sections"])
        )

    def test_catch_all_taxonomy_cannot_be_confirmed(self) -> None:
        all_papers = ", ".join(f"P{index:03d}" for index in range(1, 36))
        outline = (
            "# Review\n\n"
            "## Introduction\nSection role: introduction\nPurpose: define scope.\n\n"
            "## Other or unspecified\nSection role: body\n"
            f"Assigned papers: {all_papers}.\n\n"
            "## Conclusion\nSection role: conclusion\nPurpose: synthesize.\n"
        )
        with TestClient(self.app) as client:
            selected = self.choose_outline(client, "custom")
            saved = client.put(
                f"/api/v1/projects/{self.project_id}/planning/outline",
                json={
                    "revision": selected["matrix_revision"],
                    "outline_style": "custom",
                    "outline_md": outline,
                },
                headers=self.headers(),
            )
            self.assertEqual(200, saved.status_code, saved.text)
            generated = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint",
                json={"revision": 0},
                headers=self.headers(),
            )
            self.assertEqual(200, generated.status_code, generated.text)
            blueprint = generated.json()["section_blueprint"]
            confirmed = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint/confirm",
                json={"revision": generated.json()["blueprint_revision"]},
                headers=self.headers(),
            )
        self.assertFalse(blueprint["taxonomy_diagnostics"]["can_confirm"])
        self.assertEqual(409, confirmed.status_code, confirmed.text)
        self.assertEqual(
            "taxonomy.catch_all_body_section",
            confirmed.json()["error"]["details"]["issues"][0]["rule_id"],
        )

    def test_scope_can_be_edited_with_the_existing_outline_save(self) -> None:
        with TestClient(self.app) as client:
            selected = self.choose_outline(client, "reaction")
            planning = self.planning(client)
            scope = dict(planning["scope_contract"])
            scope["target_question"] = "Which stereocontrol strategies are transferable?"
            saved = client.put(
                f"/api/v1/projects/{self.project_id}/planning/outline",
                json={
                    "revision": selected["matrix_revision"],
                    "outline_style": "reaction",
                    "outline_md": planning["selected_outline_md"],
                    "scope_contract": scope,
                },
                headers=self.headers(),
            )
            self.assertEqual(200, saved.status_code, saved.text)
            reloaded = self.planning(client)
        self.assertEqual(
            "Which stereocontrol strategies are transferable?",
            reloaded["scope_contract"]["target_question"],
        )
        self.assertEqual("user_edited", reloaded["scope_contract"]["source"])

    def test_duplicate_body_assignment_becomes_supporting_cross_reference(self) -> None:
        with TestClient(self.app) as client:
            selected = self.choose_outline(client, "custom")
            outline = (
                "# Review\n\n"
                "## Introduction\n"
                "Section role: introduction\n"
                "Purpose: define scope.\n\n"
                "## 1. First evidence theme\n"
                "Section role: body\n"
                "Assigned papers: P001, P002.\n\n"
                "## 2. Cross-cutting theme\n"
                "Section role: body\n"
                "Assigned papers: P001, P003.\n\n"
                "## Conclusion\n"
                "Section role: conclusion\n"
                "Purpose: synthesize findings.\n"
            )
            saved = client.put(
                f"/api/v1/projects/{self.project_id}/planning/outline",
                json={
                    "revision": selected["matrix_revision"],
                    "outline_style": "custom",
                    "outline_md": outline,
                },
                headers=self.headers(),
            )
            self.assertEqual(200, saved.status_code, saved.text)
            generated = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint",
                json={"revision": 0},
                headers=self.headers(),
            )
            self.assertEqual(200, generated.status_code, generated.text)
        sections = generated.json()["section_blueprint"]["sections"]
        first = next(section for section in sections if section["title"] == "First evidence theme")
        second = next(section for section in sections if section["title"] == "Cross-cutting theme")
        self.assertEqual(["P001", "P002"], first["primary_papers"])
        self.assertEqual(["P003"], second["primary_papers"])
        self.assertEqual(["P001"], second["supporting_papers"])

    def test_blueprint_confirmation_advances_to_sections(self) -> None:
        with TestClient(self.app) as client:
            self.choose_outline(client, "reaction")
            generated = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint",
                json={"revision": 0},
                headers=self.headers(),
            )
            self.assertEqual(200, generated.status_code, generated.text)
            confirmed = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint/confirm",
                json={"revision": generated.json()["blueprint_revision"]},
                headers=self.headers(),
            )
            project = client.get(f"/api/v1/projects/{self.project_id}").json()
        self.assertEqual(200, confirmed.status_code, confirmed.text)
        self.assertEqual("sections", project["current_stage"])

    def test_previous_blueprint_can_be_restored_as_a_new_review_version(self) -> None:
        with TestClient(self.app) as client:
            self.choose_outline(client, "reaction")
            first = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint",
                json={"revision": 0},
                headers=self.headers(),
            )
            self.assertEqual(200, first.status_code, first.text)
            first_payload = first.json()
            second = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint",
                json={"revision": first_payload["blueprint_revision"]},
                headers=self.headers(),
            )
            self.assertEqual(200, second.status_code, second.text)
            restored = client.post(
                f"/api/v1/projects/{self.project_id}/planning/blueprint/restore",
                json={
                    "revision": second.json()["blueprint_revision"],
                    "artifact_id": first_payload["blueprint_artifact_id"],
                },
                headers=self.headers(),
            )
            planning = self.planning(client)

        self.assertEqual(200, restored.status_code, restored.text)
        restored_payload = restored.json()
        self.assertEqual("review", restored_payload["status"])
        self.assertNotEqual(
            first_payload["blueprint_artifact_id"],
            restored_payload["blueprint_artifact_id"],
        )
        self.assertEqual(
            first_payload["blueprint_artifact_id"],
            planning["section_blueprint"]["restructure_record"][
                "restored_from_artifact_id"
            ],
        )

    def test_planning_contract_exposes_composite_tabs(self) -> None:
        with TestClient(self.app) as client:
            payload = self.planning(client)
        self.assertEqual(["matrix", "blueprint"], [tab["id"] for tab in payload["workspace"]["tabs"]])
        self.assertEqual("文献矩阵", payload["workspace"]["tabs"][0]["labels"]["zh"])
        self.assertEqual("Blueprint", payload["workspace"]["tabs"][1]["labels"]["en"])

    def test_planning_api_and_container_are_user_isolated(self) -> None:
        self.assertIs(
            self.app.state.planning_service,
            self.app.state.container.planning_service,
        )
        self.current = self.second
        with TestClient(self.app) as client:
            response = client.get(f"/api/v1/projects/{self.project_id}/planning")
        self.assertEqual(404, response.status_code, response.text)


if __name__ == "__main__":
    unittest.main()
